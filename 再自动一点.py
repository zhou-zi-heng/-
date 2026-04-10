import streamlit as st  
import streamlit.components.v1 as components  
from openai import OpenAI  
import io, base64, json, re, requests, uuid, copy  
from datetime import datetime  
from docx import Document  
  
# ==========================================  
# 1. 页面全局配置与前端美化  
# ==========================================  
st.set_page_config(page_title="ZenMux 创作者工作站", page_icon="🐙", layout="wide")  
st.markdown("""  
    <style>  
    .stButton>button { border-radius: 8px; font-weight: bold; transition: all 0.3s; }  
    .stChatInput { padding-bottom: 20px; }  
    button[title="View fullscreen"] {display: none;}  
    .css-1jc7ptx, .e1ewe7hr3, .viewerBadge_container__1QSob, .styles_viewerBadge__1yB5_ {display: none;}  
    </style>  
""", unsafe_allow_html=True)  
  
# ==========================================  
# 2. 数据状态管理 (变更追踪 + 云端安全)  
# ==========================================  
def _track():  
    st.session_state._unsaved = st.session_state.get("_unsaved", 0) + 1  
  
def save_profiles(): _track()  
def save_sops(): _track()  
def save_memory(): _track()  
def save_free_chats(): _track()  
  
# ==========================================  
# 3. 内置 SOP 模板库  
# ==========================================  
BUILTIN_TEMPLATES = {  
    "📖 悬疑推理小说": {  
        "memory_mode": "manual",  
        "system_prompt": "你是一名顶尖悬疑推理小说家，擅长构建精密逻辑链条和出人意料的反转。文风冷峻，节奏紧凑，善用细节伏笔。绝不输出废话。",  
        "negative_memory": [],  
        "steps": [  
            {"prompt": "为主题【{主题}】构思完整悬疑故事大纲：核心诡计、人物表、线索链、章节梗概。", "loop": 1, "reference": "", "enable_word_control": False, "target_words": 0, "word_tolerance": 5, "max_corrections": 2},  
            {"prompt": "根据大纲撰写第【{循环索引}】章正文。要求：伏笔自然、节奏紧凑、结尾留悬念。", "loop": 5, "reference": "", "enable_word_control": True, "target_words": 2000, "word_tolerance": 5, "max_corrections": 2},  
            {"prompt": "撰写最终章：真相揭晓，所有伏笔收束。结尾写'全文完'。", "loop": 1, "reference": "", "enable_word_control": True, "target_words": 3000, "word_tolerance": 10, "max_corrections": 2}  
        ],  
        "triggers": [{"type": "terminate", "keyword": "全文完", "action": ""}]  
    },  
    "💕 都市言情小说": {  
        "memory_mode": "manual",  
        "system_prompt": "你是细腻的都市言情小说家，善于捕捉情感细节和生活质感。对话自然，情感真挚，拒绝狗血。",  
        "negative_memory": [],  
        "steps": [  
            {"prompt": "为主题【{主题}】设计言情大纲：核心冲突、双主角设定、感情线节点。", "loop": 1, "reference": "", "enable_word_control": False, "target_words": 0, "word_tolerance": 5, "max_corrections": 2},  
            {"prompt": "撰写第【{循环索引}】章，注意情感推进和细节。", "loop": 8, "reference": "", "enable_word_control": True, "target_words": 2500, "word_tolerance": 5, "max_corrections": 2}  
        ],  
        "triggers": [{"type": "terminate", "keyword": "全文完", "action": ""}]  
    },  
    "🎬 短视频脚本批量": {  
        "memory_mode": "manual",  
        "system_prompt": "你是爆款短视频编剧，精通情绪钩子和节奏控制。每个脚本要有强开头、高密度中段、记忆点结尾。",  
        "negative_memory": [],  
        "steps": [  
            {"prompt": "为主题【{主题}】编写第【{循环索引}】条60秒短视频脚本：画面描述、旁白、字幕、BGM建议。", "loop": 5, "reference": "", "enable_word_control": True, "target_words": 300, "word_tolerance": 15, "max_corrections": 1}  
        ],  
        "triggers": []  
    }  
}  
  
def _default_step():  
    return {"prompt": "", "loop": 1, "reference": "", "enable_word_control": False, "target_words": 0, "word_tolerance": 5, "max_corrections": 2}  
  
def _ensure_step(s):  
    d = _default_step()  
    for k, v in d.items():  
        if k not in s:  
            s[k] = v  
    return s  
  
def _ensure_chat(c):  
    if "session_knowledge" not in c:  
        c["session_knowledge"] = []  
    if "system_prompt" not in c:  
        c["system_prompt"] = ""  
    return c  
  
# ==========================================  
# 4. 核心底层辅助函数  
# ==========================================  
def render_copy_button(text):  
    b64 = base64.b64encode(text.encode("utf-8")).decode("utf-8")  
    uid = uuid.uuid4().hex[:8]  
    html = (  
        '<div style="display:flex;justify-content:flex-end;padding-right:10px;">'  
        '<button id="cb' + uid + '" onclick="(function(b){navigator.clipboard.writeText('  
        "decodeURIComponent(escape(atob('" + b64 + "')))).then(function(){"  
        "b.innerText='\\u2705 \\u5df2\\u590d\\u5236';b.style.color='#4CAF50';"  
        "setTimeout(function(){b.innerText='\\ud83d\\udccb \\u590d\\u5236';b.style.color='#aaa';},2000);"  
        '})})(this)" '  
        'style="border:none;background:transparent;color:#aaa;cursor:pointer;font-size:12px;'  
        'font-weight:bold;padding:5px 10px;border-radius:6px;">📋 复制</button></div>'  
    )  
    components.html(html, height=30)  
  
def clean_novel_text(text):  
    text = re.sub(  
        r'^\s*(好的|没问题|非常荣幸|收到|为你生成|以下是|这是为您|正文开始|下面是).*?[:：]\n*',  
        '', text, flags=re.MULTILINE | re.IGNORECASE  
    )  
    text = re.sub(r'^\s*第[零一二三四五六七八九十百千0-9]+[章回节卷].*?\n', '', text, flags=re.MULTILINE)  
    text = re.sub(r'```[a-zA-Z]*\n?', '', text)  
    text = re.sub(r'\n*(希望这|如果有需要|请告诉我|期待您的反馈).*$', '', text, flags=re.IGNORECASE)  
    text = re.sub(r'\n{3,}', '\n\n', text)  
    return text.strip()  
  
def count_words(text):  
    chinese = len(re.findall(r'[\u4e00-\u9fff]', text))  
    english = len(re.findall(r'[a-zA-Z]+', text))  
    return chinese + english  
  
def generate_word_doc(messages, is_pure=False):  
    doc = Document()  
    doc.add_heading('AI 创作者工作站生成文档', 0)  
    for msg in messages:  
        if msg["role"] == "system" or not msg.get("selected", True):  
            continue  
        if is_pure:  
            if msg["role"] == "assistant":  
                doc.add_paragraph(clean_novel_text(msg["content"]))  
        else:  
            if msg["role"] == "user":  
                doc.add_heading("📌 指令", level=2)  
                doc.add_paragraph(msg["content"])  
            elif msg["role"] == "assistant":  
                doc.add_heading("🤖 AI", level=2)  
                doc.add_paragraph(msg["content"])  
    bio = io.BytesIO()  
    doc.save(bio)  
    return bio.getvalue()  
  
def export_to_pretty_html(messages, title, meta=None):  
    """  
    meta 字典可传入：  
    {  
        "source": "自由聊天" 或 "自动化流水线",  
        "system_prompt": "...",  
        "sop_name": "...",  
        "model": "...",  
        "files": [{"filename":"xx.txt","size":12345}, ...],  
        "global_file_name": "设定集.txt",  
        "topic": "..."  
    }  
    """  
    if meta is None:  
        meta = {}  
  
    css = """  
    * { margin:0; padding:0; box-sizing:border-box; }  
    body { font-family: -apple-system, 'PingFang SC', 'Microsoft YaHei', 'Segoe UI', sans-serif; background:#f0f2f5; color:#1a1a1a; }  
    .header { background:linear-gradient(135deg,#667eea 0%,#764ba2 100%); color:#fff; padding:24px 32px; position:sticky; top:0; z-index:100; box-shadow:0 2px 12px rgba(0,0,0,.15); }  
    .header h1 { font-size:22px; font-weight:700; }  
    .header .meta { font-size:12px; opacity:.75; margin-top:6px; }  
    .chat-container { max-width:860px; margin:0 auto; padding:24px 16px 80px; }  
    .info-card { background:#fff; border-radius:12px; padding:20px 24px; margin-bottom:24px; box-shadow:0 1px 4px rgba(0,0,0,.06); border-left:4px solid #667eea; }  
    .info-card h3 { font-size:14px; color:#667eea; margin-bottom:12px; font-weight:700; }  
    .info-row { display:flex; margin-bottom:8px; font-size:13px; line-height:1.6; }  
    .info-label { color:#888; min-width:90px; flex-shrink:0; font-weight:600; }  
    .info-value { color:#333; word-break:break-all; }  
    .info-value.prompt { background:#f8f8f8; padding:8px 12px; border-radius:6px; font-size:12px; line-height:1.7; margin-top:4px; white-space:pre-wrap; max-height:200px; overflow-y:auto; }  
    .file-tag { display:inline-block; background:#f0f0f0; padding:2px 10px; border-radius:12px; font-size:12px; color:#555; margin:2px 4px 2px 0; }  
    .msg { display:flex; gap:12px; margin-bottom:24px; align-items:flex-start; }  
    .msg.user { flex-direction:row-reverse; }  
    .avatar { width:36px; height:36px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:18px; flex-shrink:0; }  
    .msg.ai .avatar { background:#e8f5e9; }  
    .msg.user .avatar { background:#e3f2fd; }  
    .bubble { max-width:75%; padding:14px 18px; border-radius:16px; line-height:1.8; font-size:15px; word-wrap:break-word; white-space:pre-wrap; box-shadow:0 1px 3px rgba(0,0,0,.06); }  
    .msg.ai .bubble { background:#fff; border-top-left-radius:4px; }  
    .msg.user .bubble { background:#d1e7ff; border-top-right-radius:4px; }  
    .bubble p { margin-bottom:10px; } .bubble p:last-child { margin-bottom:0; }  
    .bubble ul, .bubble ol { padding-left:20px; margin:8px 0; }  
    .bubble li { margin-bottom:4px; }  
    .bubble code { background:#f5f5f5; padding:2px 6px; border-radius:4px; font-size:13px; font-family:'Fira Code',Consolas,monospace; }  
    .bubble pre { background:#f5f5f5; padding:12px; border-radius:8px; overflow-x:auto; margin:10px 0; }  
    .bubble pre code { background:none; padding:0; }  
    .bubble strong, .bubble b { font-weight:600; }  
    .bubble h1,.bubble h2,.bubble h3 { font-weight:700; margin:12px 0 8px; }  
    .bubble h1 { font-size:20px; } .bubble h2 { font-size:17px; } .bubble h3 { font-size:15px; }  
    .bubble blockquote { border-left:3px solid #ccc; padding-left:12px; color:#666; margin:8px 0; }  
    .bubble table { border-collapse:collapse; margin:10px 0; font-size:14px; }  
    .bubble th, .bubble td { border:1px solid #ddd; padding:6px 10px; text-align:left; }  
    .bubble th { background:#f9f9f9; font-weight:600; }  
    .word-count { font-size:11px; color:#999; margin-top:6px; text-align:right; }  
    .msg.user .word-count { text-align:left; }  
    .footer { text-align:center; padding:32px; font-size:12px; color:#aaa; border-top:1px solid #e5e5e5; max-width:860px; margin:0 auto; }  
    .copy-btn { display:inline-block; margin-top:8px; padding:4px 12px; font-size:12px; color:#888; border:1px solid #ddd; border-radius:6px; cursor:pointer; background:#fff; transition:.2s; }  
    .copy-btn:hover { color:#4CAF50; border-color:#4CAF50; background:#f0f9f0; }  
    .toggle-info { display:inline-block; font-size:12px; color:#667eea; cursor:pointer; margin-left:12px; text-decoration:underline; }  
    @media(max-width:600px) { .bubble { max-width:88%; font-size:14px; } .header h1 { font-size:18px; } .info-row { flex-direction:column; } .info-label { min-width:auto; margin-bottom:2px; } }  
    @media print { .header{position:relative;} .copy-btn{display:none;} }  
    """  
  
    js = """  
    <script>  
    function copyText(btn, id) {  
        var el = document.getElementById(id);  
        var text = el.innerText || el.textContent;  
        navigator.clipboard.writeText(text).then(function(){  
            btn.innerText='\\u2705 \\u5df2\\u590d\\u5236';  
            btn.style.color='#4CAF50';  
            setTimeout(function(){ btn.innerText='\\ud83d\\udccb \\u590d\\u5236\\u6587\\u672c'; btn.style.color='#888'; },2000);  
        });  
    }  
    function toggleInfo() {  
        var el = document.getElementById('infoCard');  
        if(el.style.display==='none'){el.style.display='block';}else{el.style.display='none';}  
    }  
    </script>  
    """  
  
    # === 构建信息卡片 ===  
    info_html = ""  
    has_meta = any(meta.get(k) for k in ["source", "system_prompt", "sop_name", "model", "files", "global_file_name", "topic"])  
    if has_meta:  
        rows = ""  
        if meta.get("source"):  
            rows += '<div class="info-row"><span class="info-label">\U0001f4cd \u6765\u6e90</span><span class="info-value">' + meta["source"] + '</span></div>'  
        if meta.get("topic"):  
            rows += '<div class="info-row"><span class="info-label">\U0001f3af \u4e3b\u9898</span><span class="info-value">' + meta["topic"] + '</span></div>'  
        if meta.get("sop_name"):  
            rows += '<div class="info-row"><span class="info-label">\U0001f9e9 SOP</span><span class="info-value">' + meta["sop_name"] + '</span></div>'  
        if meta.get("model"):  
            rows += '<div class="info-row"><span class="info-label">\U0001f9e0 \u6a21\u578b</span><span class="info-value">' + meta["model"] + '</span></div>'  
        if meta.get("files"):  
            tags = ""  
            for f in meta["files"]:  
                size_str = "{:,}".format(f.get("size", 0))  
                tags += '<span class="file-tag">\U0001f4c4 ' + f["filename"] + ' (' + size_str + ' \u5b57)</span>'  
            rows += '<div class="info-row"><span class="info-label">\U0001f4ce \u6302\u8f7d\u6587\u4ef6</span><span class="info-value">' + tags + '</span></div>'  
        if meta.get("global_file_name"):  
            rows += '<div class="info-row"><span class="info-label">\U0001f4c1 \u8bbe\u5b9a\u96c6</span><span class="info-value">' + meta["global_file_name"] + '</span></div>'  
        if meta.get("system_prompt"):  
            import html as html_mod  
            safe_sp = html_mod.escape(meta["system_prompt"])  
            rows += '<div class="info-row"><span class="info-label">\U0001f3ad \u4eba\u8bbe</span></div><div class="info-value prompt">' + safe_sp + '</div>'  
  
        info_html = '<div class="info-card" id="infoCard">' + '<h3>\u2699\ufe0f \u5bf9\u8bdd\u914d\u7f6e\u4fe1\u606f</h3>' + rows + '</div>'  
  
    # === 构建消息 ===  
    msg_html = ""  
    msg_idx = 0  
    for m in messages:  
        if m["role"] == "system":  
            continue  
        is_user = m["role"] == "user"  
        role_class = "user" if is_user else "ai"  
        avatar = "\U0001f64b" if is_user else "\U0001f916"  
        content = m["content"]  
  
        import html as html_mod  
        safe = html_mod.escape(content)  
        safe = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', safe)  
        safe = re.sub(r'`([^`]+)`', r'<code>\1</code>', safe)  
        safe = re.sub(r'^### (.+)$', r'<h3>\1</h3>', safe, flags=re.MULTILINE)  
        safe = re.sub(r'^## (.+)$', r'<h2>\1</h2>', safe, flags=re.MULTILINE)  
        safe = re.sub(r'^# (.+)$', r'<h1>\1</h1>', safe, flags=re.MULTILINE)  
        safe = re.sub(r'^\- (.+)$', r'<li>\1</li>', safe, flags=re.MULTILINE)  
        safe = re.sub(r'((?:<li>.*</li>\n?)+)', r'<ul>\1</ul>', safe)  
        safe = re.sub(r'^\d+\.\s+(.+)$', r'<li>\1</li>', safe, flags=re.MULTILINE)  
        safe = re.sub(r'\n\n+', '</p><p>', safe)  
        safe = '<p>' + safe + '</p>'  
        safe = safe.replace('\n', '<br>')  
  
        wc = count_words(content)  
        block_id = "block_" + str(msg_idx)  
        copy_html = ""  
        if not is_user:  
            copy_html = '<button class="copy-btn" onclick="copyText(this,\'' + block_id + '\')">\U0001f4cb \u590d\u5236\u6587\u672c</button>'  
  
        msg_html += (  
            '<div class="msg ' + role_class + '">'  
            '<div class="avatar">' + avatar + '</div><div>'  
            '<div class="bubble" id="' + block_id + '">' + safe + '</div>'  
            '<div class="word-count">' + str(wc) + ' \u5b57</div>'  
            + copy_html + '</div></div>'  
        )  
        msg_idx += 1  
  
    date_str = datetime.now().strftime('%Y-%m-%d %H:%M')  
    total_ai_words = sum(count_words(m["content"]) for m in messages if m["role"] == "assistant")  
    total_msgs = sum(1 for m in messages if m["role"] != "system")  
  
    toggle_link = ""  
    if has_meta:  
        toggle_link = '<span class="toggle-info" onclick="toggleInfo()">\u5c55\u5f00/\u6536\u8d77\u914d\u7f6e\u4fe1\u606f</span>'  
  
    full_html = (  
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">'  
        '<meta name="viewport" content="width=device-width,initial-scale=1">'  
        '<title>' + title + '</title><style>' + css + '</style></head><body>'  
        '<div class="header"><h1>\U0001f4ac ' + title + '</h1>'  
        '<div class="meta">' + date_str + ' | '  
        + str(total_msgs) + ' \u6761\u6d88\u606f | AI \u5171 '  
        + '{:,}'.format(total_ai_words) + ' \u5b57'  
        + toggle_link + '</div></div>'  
        '<div class="chat-container">' + info_html + msg_html + '</div>'  
        '<div class="footer">ZenMux AI \u521b\u4f5c\u8005\u5de5\u4f5c\u7ad9</div>'  
        + js + '</body></html>'  
    )  
    return full_html.encode('utf-8')  


  
def fetch_models(base_url, api_key):  
    try:  
        url = (base_url.strip().rstrip('/') or "https://api.openai.com/v1") + "/models"  
        resp = requests.get(url, headers={"Authorization": "Bearer " + api_key.strip()}, timeout=8)  
        if resp.status_code == 200:  
            return True, sorted([m["id"] for m in resp.json().get("data", [])])  
        else:  
            return False, "状态码 " + str(resp.status_code) + ": " + resp.text[:100]  
    except Exception as e:  
        return False, str(e)  
  
def get_client():  
    p = st.session_state.profiles[st.session_state.active_profile_idx]  
    url = p["base_url"].strip() or "https://api.openai.com/v1"  
    return OpenAI(base_url=url, api_key=p["api_key"].strip()), p  
  
def build_api_kwargs(profile, api_msgs):  
    kw = {"model": profile["model"], "messages": api_msgs, "stream": True}  
    if profile.get("use_temperature", True):  
        kw["temperature"] = profile.get("temperature", 0.7)  
    if profile.get("use_max_tokens", True):  
        kw["max_tokens"] = profile.get("max_tokens", 4096)  
    if profile.get("use_top_p", False):  
        kw["top_p"] = profile.get("top_p", 1.0)  
    if profile.get("use_frequency_penalty", False):  
        kw["frequency_penalty"] = profile.get("frequency_penalty", 0.0)  
    return kw  
  
def stream_generator(api_stream):  
    st.session_state.auto_engine["last_finish_reason"] = "stop"  
    for chunk in api_stream:  
        if chunk.choices and chunk.choices[0].delta.content is not None:  
            yield chunk.choices[0].delta.content  
        if chunk.choices and chunk.choices[0].finish_reason is not None:  
            st.session_state.auto_engine["last_finish_reason"] = chunk.choices[0].finish_reason  
  
# ==========================================  
# 5. 状态初始化 (完整私有沙盒)  
# ==========================================  
if "initialized" not in st.session_state:  
    st.session_state.profiles = [{  
        "name": "默认引擎", "base_url": "", "api_key": "",  
        "model": "anthropic/claude-sonnet-4.6",  
        "use_temperature": True, "temperature": 0.8,  
        "use_max_tokens": True, "max_tokens": 4096,  
        "use_top_p": False, "top_p": 1.0,  
        "use_frequency_penalty": False, "frequency_penalty": 0.0  
    }]  
    default_step = _default_step()  
    default_step["prompt"] = "撰写第【{循环索引}】章"  
    default_step["loop"] = 2  
    st.session_state.sops = {  
        "小说账号预设": {  
            "memory_mode": "manual",  
            "system_prompt": "你是一名悬疑小说家，文风冷峻，绝不输出废话。",  
            "negative_memory": [],  
            "steps": [default_step],  
            "triggers": [{"type": "terminate", "keyword": "全文完", "action": ""}]  
        }  
    }  
    st.session_state.memory = {}  
    first_id = str(uuid.uuid4())  
    st.session_state.free_chats = {  
        first_id: {"title": "新对话", "messages": [], "session_knowledge": [], "system_prompt": ""}  
    }  
    st.session_state.active_profile_idx = 0  
    st.session_state.current_page = "🤖 自动化流水线"  
    st.session_state.current_chat_id = first_id  
    st.session_state.auto_engine = {  
        "is_running": False, "is_paused": False, "is_finished": False, "messages": [],  
        "sop_name": "", "topic": "", "global_file": "",  
        "current_step_idx": 0, "current_loop_idx": 1,  
        "pending_instruction": "", "last_finish_reason": "",  
        "correction_count": 0, "word_count_log": [], "model_bias_history": []  
    }  
    st.session_state._unsaved = 0  
    st.session_state.initialized = True  
  
# ==========================================  
# 6. 全局侧边栏导航  
# ==========================================  
with st.sidebar:  
    col_img, col_txt = st.columns([1, 3])  
    with col_img:  
        st.image("https://api.iconify.design/fluent-emoji:octopus.svg?width=80", width=45)  
    with col_txt:  
        st.header("控制中枢")  
  
    if st.session_state.get("_unsaved", 0) > 3:  
        st.warning("⚠️ 有 " + str(st.session_state._unsaved) + " 项未备份更改，建议导出快照！")  
  
    st.write("")  
    pages = ["🤖 自动化流水线", "💬 自由聊天区", "📝 账号SOP与灵魂", "⚙️ 底层引擎配置"]  
    for pg in pages:  
        btype = "primary" if st.session_state.current_page == pg else "secondary"  
        if st.button(pg, use_container_width=True, type=btype):  
            st.session_state.current_page = pg  
            st.rerun()  
  
    active_p = st.session_state.profiles[st.session_state.active_profile_idx]  
    st.divider()  
    st.caption("🟢 **当前挂载**: " + active_p['name'] + "\n🧠 **模型**: " + active_p['model'])  
  
    # === 自由聊天区侧边栏 ===  
    if st.session_state.current_page == "💬 自由聊天区":  
        st.divider()  
        st.header("📚 历史对话")  
        if st.button("➕ 开启新对话", use_container_width=True, type="primary"):  
            nid = str(uuid.uuid4())  
            st.session_state.free_chats[nid] = {  
                "title": "新对话", "messages": [], "session_knowledge": [], "system_prompt": ""  
            }  
            st.session_state.current_chat_id = nid  
            save_free_chats()  
            st.rerun()  
  
        for c_id, c_data in reversed(list(st.session_state.free_chats.items())):  
            _ensure_chat(c_data)  
            t = c_data["title"]  
            c_title = t[:12] + "..." if len(t) > 12 else t  
            lbl = ("⭐ " if c_id == st.session_state.current_chat_id else "📄 ") + c_title  
            if st.button(lbl, key="chat_" + c_id, use_container_width=True):  
                st.session_state.current_chat_id = c_id  
                st.rerun()  
  
        st.divider()  
        curr_msgs_sb = st.session_state.free_chats[st.session_state.current_chat_id]["messages"]  
        if curr_msgs_sb:  
            st.markdown("**📦 导出当前对话**")  
            exp_mode = st.radio("格式", ["完整记录", "纯享正文"], label_visibility="collapsed")  
            is_pure = (exp_mode == "纯享正文")  
            if is_pure:  
                txt_c = "\n\n".join([clean_novel_text(m['content']) for m in curr_msgs_sb if m['role'] == 'assistant'])  
            else:  
                parts = []  
                for m in curr_msgs_sb:  
                    role_name = "我" if m['role'] == 'user' else "AI"  
                    parts.append(role_name + ":\n" + m['content'] + "\n\n" + "-" * 40 + "\n")  
                txt_c = "\n".join(parts)  
  
            chat_title_for_export = st.session_state.free_chats[st.session_state.current_chat_id]["title"]  
            chat_data_for_export = st.session_state.free_chats[st.session_state.current_chat_id]  
            _ensure_chat(chat_data_for_export)  
            export_meta_chat = {  
                "source": "自由聊天区",  
                "system_prompt": chat_data_for_export.get("system_prompt", ""),  
                "model": active_p["model"],  
                "files": [  
                    {"filename": k["filename"], "size": len(k["content"])}  
                    for k in chat_data_for_export.get("session_knowledge", [])  
                ]  
            }  
  
            ec1, ec2, ec3 = st.columns(3)  
            with ec1:  
                st.download_button(  
                    "📥 TXT", txt_c.encode('utf-8'), "对话.txt",  
                    "text/plain", use_container_width=True  
                )  
            with ec2:  
                st.download_button(  
                    "📥 Word", generate_word_doc(curr_msgs_sb, is_pure), "对话.docx",  
                    use_container_width=True  
                )  
            with ec3:  
                st.download_button(  
                    "🎨 HTML",  
                    export_to_pretty_html(curr_msgs_sb, chat_title_for_export, export_meta_chat),  
                    "对话.html", "text/html", use_container_width=True  
                )  
  
        # 删除对话（二次确认）  
        if st.session_state.get("_confirm_del_chat"):  
            st.error("确认删除此对话？不可撤回！")  
            dcc1, dcc2 = st.columns(2)  
            if dcc1.button("✅ 确认", key="yes_del_chat"):  
                if len(st.session_state.free_chats) > 1:  
                    del st.session_state.free_chats[st.session_state.current_chat_id]  
                else:  
                    st.session_state.free_chats[st.session_state.current_chat_id] = {  
                        "title": "新对话", "messages": [], "session_knowledge": [], "system_prompt": ""  
                    }  
                st.session_state.current_chat_id = list(st.session_state.free_chats.keys())[-1]  
                st.session_state._confirm_del_chat = False  
                save_free_chats()  
                st.rerun()  
            if dcc2.button("❌ 取消", key="no_del_chat"):  
                st.session_state._confirm_del_chat = False  
                st.rerun()  
        else:  
            if st.button("🗑️ 删除对话", use_container_width=True):  
                st.session_state._confirm_del_chat = True  
                st.rerun()  
  
    # === 全量资产导出恢复舱 ===  
    st.divider()  
    with st.expander("📦 全量资产导出恢复舱", expanded=False):  
        st.caption("换电脑时一键导入导出所有数据。")  
        full_data = json.dumps({  
            "profiles": st.session_state.profiles,  
            "sops": st.session_state.sops,  
            "memory": st.session_state.memory,  
            "free_chats": st.session_state.free_chats  
        }, ensure_ascii=False, indent=2).encode('utf-8')  
        fname = "ZenMux_Backup_" + datetime.now().strftime('%m%d_%H%M') + ".json"  
        if st.download_button("📥 导出全量快照包", full_data, fname, "application/json", use_container_width=True, type="primary"):  
            st.session_state._unsaved = 0  
  
        uploaded_ws = st.file_uploader("📂 导入快照 (覆盖当前)", type="json")  
        if uploaded_ws:  
            try:  
                data = json.loads(uploaded_ws.getvalue().decode('utf-8'))  
                st.session_state.profiles = data.get("profiles", st.session_state.profiles)  
                st.session_state.sops = data.get("sops", st.session_state.sops)  
                for sop in st.session_state.sops.values():  
                    for s in sop.get("steps", []):  
                        _ensure_step(s)  
                st.session_state.memory = data.get("memory", st.session_state.memory)  
                st.session_state.free_chats = data.get("free_chats", st.session_state.free_chats)  
                for ch in st.session_state.free_chats.values():  
                    _ensure_chat(ch)  
                if not st.session_state.free_chats:  
                    nid = str(uuid.uuid4())  
                    st.session_state.free_chats = {  
                        nid: {"title": "新对话", "messages": [], "session_knowledge": [], "system_prompt": ""}  
                    }  
                st.session_state.active_profile_idx = 0  
                st.session_state.current_chat_id = list(st.session_state.free_chats.keys())[-1]  
                st.session_state._unsaved = 0  
                st.success("✅ 恢复成功！")  
                st.rerun()  
            except Exception as e:  
                st.error("导入失败: " + str(e))  

  
# ==========================================  
# 模块 1: 自动化流水线  
# ==========================================  
if st.session_state.current_page == "🤖 自动化流水线":  
    engine = st.session_state.auto_engine  
    col_ctrl, col_view = st.columns([1, 2.5])  
  
    with col_ctrl:  
        st.header("⚙️ 调度控制台")  
  
        if engine["is_running"]:  
            if engine.get("is_paused"):  
                st.info("⏸️ 引擎已暂停")  
                pause_input = st.text_area(  
                    "💬 注入修正指令（留空直接继续）", "",  
                    placeholder="例如：接下来减少对话，多写心理描写", height=80  
                )  
                cp1, cp2 = st.columns(2)  
                if cp1.button("▶️ 继续执行", type="primary", use_container_width=True):  
                    if pause_input.strip():  
                        engine["pending_instruction"] = pause_input.strip()  
                    engine["is_paused"] = False  
                    st.rerun()  
                if cp2.button("⏹️ 彻底停止", use_container_width=True):  
                    engine["is_running"] = False  
                    engine["is_paused"] = False  
                    st.rerun()  
            else:  
                st.warning("⚠️ 引擎高速运转中...")  
                sop_steps_total = max(len(st.session_state.sops.get(engine["sop_name"], {"steps": []})["steps"]), 1)  
                st.progress(min(engine["current_step_idx"] / sop_steps_total, 1.0))  
                info_str = (  
                    "阶段 " + str(engine['current_step_idx'] + 1) + "/" + str(sop_steps_total) +  
                    " | 循环 " + str(engine['current_loop_idx']) +  
                    " | 矫正第" + str(engine['correction_count']) + "轮"  
                )  
                st.caption(info_str)  
  
                ai_msgs_running = [m for m in engine["messages"] if m["role"] == "assistant" and m.get("selected", True)]  
                if ai_msgs_running:  
                    total_wc_run = sum(count_words(m["content"]) for m in ai_msgs_running)  
                    st.metric("📊 已生成总字数", "{:,}".format(total_wc_run))  
  
                cp1, cp2 = st.columns(2)  
                if cp1.button("⏸️ 暂停", use_container_width=True):  
                    engine["is_paused"] = True  
                    st.rerun()  
                if cp2.button("⏹️ 强制急停", type="primary", use_container_width=True):  
                    engine["is_running"] = False  
                    st.rerun()  
        else:  
            if not st.session_state.sops:  
                st.warning("请先去配置一个 SOP。")  
                st.stop()  
            sel_sop = st.selectbox("1. 挂载执行 SOP (账号人设)", list(st.session_state.sops.keys()))  
            in_topic = st.text_input("2. 注入 {主题}", placeholder="例如：赛博朋克修仙传")  
            up_file = st.file_uploader("3. 挂载全局设定集 (可选)", type=['txt', 'md'])  
  
            if st.button("🚀 点火启动", type="primary", use_container_width=True):  
                if not active_p["api_key"]:  
                    st.error("引擎未配置 API Key！")  
                elif not in_topic:  
                    st.error("请填入主题！")  
                else:  
                    gf = ""  
                    if up_file:  
                        gf = up_file.getvalue().decode("utf-8")  
                    engine.update({  
                        "is_running": True, "is_paused": False, "is_finished": False, "messages": [],  
                        "sop_name": sel_sop, "topic": in_topic, "global_file": gf,  
                        "current_step_idx": 0, "current_loop_idx": 1,  
                        "pending_instruction": "", "last_finish_reason": "",  
                        "correction_count": 0, "word_count_log": [], "model_bias_history": []  
                    })  
                    st.rerun()  
  
        # === 成果验收区 ===  
        st.divider()  
        if engine["messages"]:  
            st.markdown("### 📦 成果验收与记忆管理")  
            sel_msgs = [m for m in engine["messages"] if m["role"] == "assistant" and m.get("selected", True)]  
            raw_text = "\n\n".join([m["content"] for m in sel_msgs])  
            pure_text = clean_novel_text(raw_text)  
            total_export_wc = count_words(pure_text)  
            st.info("📊 选中导出内容：**{:,}** 字 | **{}** 段".format(total_export_wc, len(sel_msgs)))  
  
            if not engine["is_running"]:  
                sop_data = st.session_state.sops.get(engine["sop_name"], {})  
                mem_mode = sop_data.get("memory_mode", "manual")  
  
                if mem_mode == "manual":  
                    st.caption("🧠 记忆模式：手动提取蒸馏")  
                    if st.button("💾 存入该账号记忆保险库", type="primary", use_container_width=True):  
                        sn = engine["sop_name"]  
                        if sn not in st.session_state.memory:  
                            st.session_state.memory[sn] = []  
                        st.session_state.memory[sn].append({  
                            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),  
                            "topic": engine["topic"],  
                            "content": pure_text[:2500]  
                        })  
                        save_memory()  
                        st.toast("已安全落盘！", icon="💾")  
                else:  
                    st.caption("🌱 记忆模式：动态活体进化")  
                    feedback = st.text_input("💬 避坑反馈：", placeholder="例如：不准写废话")  
                    if st.button("提交反馈写入潜意识", use_container_width=True):  
                        if feedback.strip():  
                            if "negative_memory" not in sop_data:  
                                sop_data["negative_memory"] = []  
                            sop_data["negative_memory"].append(feedback)  
                            save_sops()  
                            if len(sop_data["negative_memory"]) >= 3:  
                                with st.spinner("融合新规则中..."):  
                                    try:  
                                        client, profile = get_client()  
                                        fusion_p = (  
                                            "原人设：" + sop_data['system_prompt'] +  
                                            "。避坑反馈：" + '; '.join(sop_data['negative_memory']) +  
                                            "。请深度融合进原人设，输出纯文本新 System Prompt。"  
                                        )  
                                        resp = client.chat.completions.create(  
                                            model=profile["model"],  
                                            messages=[{"role": "user", "content": fusion_p}]  
                                        )  
                                        sop_data["system_prompt"] = resp.choices[0].message.content.strip()  
                                        sop_data["negative_memory"] = []  
                                        save_sops()  
                                        st.success("人设已进化！")  
                                    except Exception as e:  
                                        st.error("融合失败: " + str(e))  
                            else:  
                                st.toast("反馈已记录！", icon="✅")  
  
            ec1, ec2, ec3 = st.columns(3)  
            with ec1:  
                st.download_button(  
                    "📥 全文TXT", raw_text.encode('utf-8'),  
                    engine['topic'] + "_完整.txt", use_container_width=True  
                )  
            with ec2:  
                st.download_button(  
                    "✨ 纯享TXT", pure_text.encode('utf-8'),  
                    engine['topic'] + "_纯正文.txt", use_container_width=True  
                )  
                        with ec3:  
                _sop_for_export = st.session_state.sops.get(engine["sop_name"], {})  
                _export_meta_auto = {  
                    "source": "自动化流水线",  
                    "topic": engine["topic"],  
                    "sop_name": engine["sop_name"],  
                    "system_prompt": _sop_for_export.get("system_prompt", ""),  
                    "model": active_p["model"],  
                    "global_file_name": "(已挂载)" if engine.get("global_file") else ""  
                }  
                st.download_button(  
                    "🎨 精美HTML",  
                    export_to_pretty_html(engine["messages"], engine["topic"], _export_meta_auto),  
                    engine['topic'] + ".html", "text/html", use_container_width=True  
                )  


  
            # 字数精控报告  
            if engine["word_count_log"]:  
                with st.expander("📊 字数精控报告", expanded=False):  
                    total_actual = sum(w["actual"] for w in engine["word_count_log"])  
                    total_target = sum(w["target"] for w in engine["word_count_log"])  
                    if engine["word_count_log"]:  
                        avg_dev = sum(  
                            abs(w["actual"] - w["target"]) / max(w["target"], 1)  
                            for w in engine["word_count_log"]  
                        ) / len(engine["word_count_log"])  
                    else:  
                        avg_dev = 0  
                    st.markdown(  
                        "**合计**: {:,} / 目标 {:,} | **平均偏差**: {:.1%}".format(  
                            total_actual, total_target, avg_dev  
                        )  
                    )  
                    header = "| 章节 | 目标 | 实际 | 偏差 | 矫正 |\n|---|---|---|---|---|\n"  
                    rows = ""  
                    for w in engine["word_count_log"]:  
                        dev = (w["actual"] - w["target"]) / max(w["target"], 1)  
                        emoji = "✅" if abs(dev) <= 0.05 else "⚠️"  
                        rows += (  
                            "| 阶段" + str(w['step'] + 1) + "-第" + str(w['loop']) + "轮"  
                            + " | " + str(w['target'])  
                            + " | " + str(w['actual'])  
                            + " | " + "{:+.1%}".format(dev) + " " + emoji  
                            + " | " + str(w['corrections']) + "次 |\n"  
                        )  
                    st.markdown(header + rows)  
  
            # 清理工作台（二次确认）  
            if st.session_state.get("_confirm_clear"):  
                st.error("确认清理？所有生成内容将丢失！")  
                clc1, clc2 = st.columns(2)  
                if clc1.button("✅ 确认清理", key="yes_clear"):  
                    engine.update({  
                        "messages": [], "is_finished": False, "is_running": False,  
                        "word_count_log": [], "model_bias_history": []  
                    })  
                    st.session_state._confirm_clear = False  
                    st.rerun()  
                if clc2.button("❌ 取消", key="no_clear"):  
                    st.session_state._confirm_clear = False  
                    st.rerun()  
            else:  
                if st.button("🧹 清理工作台", use_container_width=True):  
                    st.session_state._confirm_clear = True  
                    st.rerun()  
  
    # === 监视大屏 ===  
    with col_view:  
        st.header("🖥️ 监视大屏")  
        with st.container(height=750, border=True):  
            for i, msg in enumerate(engine["messages"]):  
                if msg["role"] == "system":  
                    continue  
                with st.chat_message(msg["role"]):  
                    st.markdown(msg["content"])  
                    if msg["role"] == "assistant":  
                        render_copy_button(msg["content"])  
                        ws = msg.get("_word_status", "")  
                        if ws:  
                            st.caption(ws)  
                        else:  
                            st.caption("📊 " + str(count_words(msg["content"])) + " 字")  
                        msg["selected"] = st.checkbox(  
                            "☑️ 选中导出", msg.get("selected", True), key="ac_" + str(i)  
                        )  
  
            # === 核心执行引擎 ===  
            if engine["is_running"] and not engine.get("is_paused", False):  
                client, profile = get_client()  
                sop_data = st.session_state.sops[engine["sop_name"]]  
                steps = sop_data["steps"]  
                for s in steps:  
                    _ensure_step(s)  
                triggers = sop_data.get("triggers", [])  
                curr_step = steps[engine["current_step_idx"]]  
  
                current_prompt = engine["pending_instruction"]  
                if not current_prompt:  
                    current_prompt = curr_step["prompt"].replace(  
                        "{主题}", engine["topic"]  
                    ).replace(  
                        "{循环索引}", str(engine["current_loop_idx"])  
                    )  
                engine["pending_instruction"] = ""  
  
                # 字数精控：预注入  
                wc_enabled = curr_step.get("enable_word_control", False) and curr_step.get("target_words", 0) > 0  
                word_injection = ""  
                if wc_enabled and engine["correction_count"] == 0:  
                    target_adj = curr_step["target_words"]  
                    if len(engine["model_bias_history"]) >= 2:  
                        recent = engine["model_bias_history"][-5:]  
                        bias = sum(recent) / len(recent)  
                        if abs(bias - 1.0) > 0.05:  
                            target_adj = int(curr_step["target_words"] / bias)  
                    word_injection = "\n\n【本章字数要求：严格控制在 " + str(target_adj) + " 字左右】"  
  
                silence = '\n\n【系统强制指令：不要重复上文，不准说好的、收到等废话，不准带章节标题，直接从正文第一个字开始！】'  
                final_prompt = current_prompt + word_injection + silence  
  
                engine["messages"].append({"role": "user", "content": current_prompt, "selected": False})  
                with st.chat_message("user"):  
                    st.markdown("*(⚡ 指令)*: " + current_prompt)  
  
                # 构建 API 消息包（知识置顶，最大化缓存命中）  
                api_msgs = []  
                sys_prompt = sop_data.get("system_prompt", "").strip()  
                if sys_prompt:  
                    api_msgs.append({"role": "system", "content": sys_prompt})  
                if sop_data.get("memory_mode") == "dynamic" and sop_data.get("negative_memory"):  
                    api_msgs.append({  
                        "role": "system",  
                        "content": "【避坑铁律】：" + '; '.join(sop_data['negative_memory'])  
                    })  
                if engine["global_file"]:  
                    api_msgs.append({"role": "system", "content": "【全局设定】\n" + engine['global_file']})  
                if curr_step.get("reference"):  
                    api_msgs.append({"role": "system", "content": "【本阶段设定】\n" + curr_step['reference']})  
  
                for idx, m in enumerate(engine["messages"]):  
                    if idx == len(engine["messages"]) - 1 and m["role"] == "user":  
                        api_msgs.append({"role": "user", "content": final_prompt})  
                    else:  
                        api_msgs.append({"role": m["role"], "content": m["content"]})  
  
                with st.chat_message("assistant"):  
                    try:  
                        resp = client.chat.completions.create(**build_api_kwargs(profile, api_msgs))  
                        full_resp = st.write_stream(stream_generator(resp))  
                        render_copy_button(full_resp)  
  
                        wc_actual = count_words(full_resp)  
                        msg_data = {  
                            "role": "assistant", "content": full_resp,  
                            "selected": True,  
                            "_word_status": "📊 " + str(wc_actual) + " 字"  
                        }  
                        engine["messages"].append(msg_data)  
  
                        hit_trigger = False  
  
                        # === 字数精控闸门 ===  
                        if wc_enabled:  
                            target = curr_step["target_words"]  
                            tol = curr_step.get("word_tolerance", 5) / 100.0  
                            deviation = (wc_actual - target) / max(target, 1)  
  
                            if abs(deviation) > tol:  
                                max_corr = curr_step.get("max_corrections", 2)  
                                if engine["correction_count"] < max_corr:  
                                    engine["correction_count"] += 1  
                                    msg_data["selected"] = False  
                                    msg_data["_word_status"] = (  
                                        "⚠️ " + str(wc_actual) + "字 (偏差" +  
                                        "{:+.1%}".format(deviation) + "，矫正第" +  
                                        str(engine['correction_count']) + "轮)"  
                                    )  
                                    if deviation > 0:  
                                        engine["pending_instruction"] = (  
                                            "【字数矫正】你刚才写了约" + str(wc_actual) +  
                                            "字，超出目标" + str(target) +  
                                            "字。请精简至" + str(target) +  
                                            "字以内，保持情节完整，删减冗余描写。直接输出修改后的完整正文。"  
                                        )  
                                    else:  
                                        engine["pending_instruction"] = (  
                                            "【字数矫正】你刚才只写了约" + str(wc_actual) +  
                                            "字，不足目标" + str(target) +  
                                            "字。请扩写至" + str(target) +  
                                            "字左右，增加环境描写和细节。直接输出修改后的完整正文。"  
                                        )  
                                    hit_trigger = True  
                                else:  
                                    msg_data["_word_status"] = (  
                                        "⚠️ " + str(wc_actual) + "字 (偏差" +  
                                        "{:+.1%}".format(deviation) + "，已达最大矫正次数)"  
                                    )  
                            else:  
                                msg_data["_word_status"] = (  
                                    "✅ " + str(wc_actual) + "字 (目标" + str(target) + "字)"  
                                )  
  
                        # === length 中断续写 ===  
                        if not hit_trigger and engine["last_finish_reason"] == "length":  
                            engine["pending_instruction"] = "⚠️ 因字数限制中断，请紧接上文最后一个字继续往下写。"  
                            hit_trigger = True  
  
                        # === 触发器检测 ===  
                        if not hit_trigger:  
                            for t in triggers:  
                                if t.get("keyword") and t["keyword"] in full_resp:  
                                    if t["type"] == "terminate":  
                                        engine["is_running"] = False  
                                        engine["is_finished"] = True  
                                        hit_trigger = True  
                                        break  
                                    elif t["type"] == "intervene":  
                                        engine["pending_instruction"] = t["action"]  
                                        hit_trigger = True  
                                        break  
  
                        # === 正常推进 ===  
                        if not hit_trigger:  
                            if wc_enabled:  
                                engine["word_count_log"].append({  
                                    "step": engine["current_step_idx"],  
                                    "loop": engine["current_loop_idx"],  
                                    "target": curr_step["target_words"],  
                                    "actual": wc_actual,  
                                    "corrections": engine["correction_count"]  
                                })  
                                engine["model_bias_history"].append(  
                                    wc_actual / max(curr_step["target_words"], 1)  
                                )  
                            engine["correction_count"] = 0  
  
                            if engine["current_loop_idx"] < curr_step.get("loop", 1):  
                                engine["current_loop_idx"] += 1  
                            else:  
                                engine["current_step_idx"] += 1  
                                engine["current_loop_idx"] = 1  
                            if engine["current_step_idx"] >= len(steps):  
                                engine["is_running"] = False  
                                engine["is_finished"] = True  
                        st.rerun()  
                    except Exception as e:  
                        st.error("引擎故障: " + str(e))  
                        engine["is_running"] = False  
  
# ==========================================  
# 模块 2: 自由聊天区 (知识库常驻 + 重新生成 + 编辑)  
# ==========================================  
elif st.session_state.current_page == "💬 自由聊天区":  
    curr_chat = _ensure_chat(st.session_state.free_chats[st.session_state.current_chat_id])  
    st.title("💬 " + curr_chat['title'])  
  
    # === 对话设置区 ===  
    has_content = bool(curr_chat["session_knowledge"] or curr_chat["system_prompt"])  
    with st.expander("⚙️ 对话设置（人设 / 知识库）", expanded=has_content):  
        sop_keys = list(st.session_state.sops.keys())  
        sop_choice = st.selectbox("💡 快速挂载已有 SOP 人设", ["(自定义)"] + sop_keys)  
        if sop_choice != "(自定义)":  
            prefill = st.session_state.sops[sop_choice].get("system_prompt", "")  
        else:  
            prefill = curr_chat.get("system_prompt", "")  
        curr_chat["system_prompt"] = st.text_area(  
            "🎭 System Prompt (可选)", prefill, height=80,  
            placeholder="给 AI 设定角色和规则"  
        )  
  
        up_f = st.file_uploader(  
            "📎 上传参考文件（常驻本对话，触发 API 缓存）",  
            type=['txt', 'md'],  
            key="kb_" + st.session_state.current_chat_id  
        )  
        if up_f:  
            content = up_f.getvalue().decode('utf-8')  
            already = any(k["filename"] == up_f.name for k in curr_chat["session_knowledge"])  
            if not already:  
                curr_chat["session_knowledge"].append({"filename": up_f.name, "content": content})  
                st.toast("✅ 已挂载: " + up_f.name, icon="📎")  
                st.rerun()  
  
        if curr_chat["session_knowledge"]:  
            st.markdown("**已挂载文件：**")  
            for ki, k in enumerate(curr_chat["session_knowledge"]):  
                kc1, kc2 = st.columns([4, 1])  
                kc1.caption("📄 " + k['filename'] + " (" + "{:,}".format(len(k['content'])) + " 字)")  
                if kc2.button("❌", key="rm_kb_" + str(ki)):  
                    curr_chat["session_knowledge"].pop(ki)  
                    st.rerun()  
  
    # === 聊天消息展示 ===  
    with st.container(height=600, border=False):  
        editing_idx = st.session_state.get("_editing_chat_idx")  
  
        for i, msg in enumerate(curr_chat["messages"]):  
            if msg["role"] == "system":  
                continue  
  
            with st.chat_message(msg["role"]):  
                if msg["role"] == "user" and editing_idx == i:  
                    new_text = st.text_area(  
                        "✏️ 编辑消息", msg["content"],  
                        key="edit_area_" + str(i), height=100  
                    )  
                    ebc1, ebc2 = st.columns(2)  
                    if ebc1.button("✅ 确认并重新发送", key="edit_ok_" + str(i), type="primary"):  
                        msg["content"] = new_text  
                        curr_chat["messages"] = curr_chat["messages"][:i + 1]  
                        st.session_state._editing_chat_idx = None  
                        st.session_state._auto_resend = True  
                        save_free_chats()  
                        st.rerun()  
                    if ebc2.button("❌ 取消", key="edit_cancel_" + str(i)):  
                        st.session_state._editing_chat_idx = None  
                        st.rerun()  
                else:  
                    st.markdown(msg["content"])  
                    if msg["role"] == "user" and editing_idx is None:  
                        if st.button("✏️", key="edit_btn_" + str(i)):  
                            st.session_state._editing_chat_idx = i  
                            st.rerun()  
                    if msg["role"] == "assistant":  
                        render_copy_button(msg["content"])  
                        st.caption("📊 " + str(count_words(msg["content"])) + " 字")  
                        if st.button("🔄 重新生成", key="regen_" + str(i)):  
                            curr_chat["messages"] = curr_chat["messages"][:i]  
                            st.session_state._auto_resend = True  
                            save_free_chats()  
                            st.rerun()  
  
    # === 发送逻辑 ===  
    need_resend = st.session_state.pop("_auto_resend", False)  
    prompt = st.chat_input("输入消息...")  
  
    if prompt or need_resend:  
        if not active_p["api_key"]:  
            st.error("请先配置 API Key！")  
            st.stop()  
  
        if prompt:  
            if len(curr_chat["messages"]) == 0:  
                curr_chat["title"] = prompt[:10] + "..."  
            curr_chat["messages"].append({"role": "user", "content": prompt})  
  
        # 构建 API 消息包（知识置顶，最大化缓存命中）  
        api_msgs = []  
  
        # A. System Prompt  
        sp = curr_chat.get("system_prompt", "").strip()  
        if sp:  
            api_msgs.append({"role": "system", "content": sp})  
  
        # B. 知识库全文（固定头部，触发缓存）  
        if curr_chat["session_knowledge"]:  
            kb_parts = []  
            for k in curr_chat["session_knowledge"]:  
                kb_parts.append("--- 文件: " + k['filename'] + " ---\n" + k['content'])  
            kb_text = "【以下是当前对话的核心参考文件，回答时务必参考】：\n\n" + "\n\n".join(kb_parts)  
            api_msgs.append({"role": "system", "content": kb_text})  
  
        # C. 历史消息  
        for m in curr_chat["messages"]:  
            api_msgs.append({"role": m["role"], "content": m["content"]})  
  
        client, profile = get_client()  
        with st.chat_message("assistant"):  
            try:  
                resp = client.chat.completions.create(**build_api_kwargs(profile, api_msgs))  
                full_resp = st.write_stream(resp)  
                render_copy_button(full_resp)  
                curr_chat["messages"].append({"role": "assistant", "content": full_resp})  
                save_free_chats()  
                st.rerun()  
            except Exception as e:  
                st.error("请求失败: " + str(e))  
  
# ==========================================  
# 模块 3: SOP与灵魂  
# ==========================================  
elif st.session_state.current_page == "📝 账号SOP与灵魂":  
    tab_sop, tab_vault = st.tabs(["🧩 SOP配置与人设管理", "🗄️ 账号记忆保险库"])  
  
    with tab_sop:  
        col1, col2 = st.columns([1, 2.5])  
        with col1:  
            st.subheader("账号 SOP 库")  
            sop_list = list(st.session_state.sops.keys())  
            s_name = None  
            if sop_list:  
                s_name = st.radio("选择编辑对象", sop_list)  
  
            st.divider()  
            if st.button("➕ 创建空白 SOP", use_container_width=True):  
                new_sop_name = "新账号 " + str(len(st.session_state.sops))  
                st.session_state.sops[new_sop_name] = {  
                    "memory_mode": "manual", "system_prompt": "", "negative_memory": [],  
                    "steps": [_default_step()], "triggers": []  
                }  
                save_sops()  
                st.rerun()  
  
            # 从模板创建  
            with st.expander("📦 从预设模板创建"):  
                for tpl_name, tpl_data in BUILTIN_TEMPLATES.items():  
                    if st.button("使用 " + tpl_name, key="tpl_" + tpl_name, use_container_width=True):  
                        st.session_state.sops[tpl_name] = copy.deepcopy(tpl_data)  
                        save_sops()  
                        st.rerun()  
  
            # 复制 SOP  
            if s_name:  
                if st.button("📋 复制当前 SOP", use_container_width=True):  
                    new_n = s_name + " (副本)"  
                    st.session_state.sops[new_n] = copy.deepcopy(st.session_state.sops[s_name])  
                    save_sops()  
                    st.rerun()  
  
        with col2:  
            if s_name:  
                sop = st.session_state.sops[s_name]  
  
                ca, cb = st.columns([3, 1])  
                with ca:  
                    new_name = st.text_input("✏️ 账号名称", s_name)  
                with cb:  
                    st.write("")  
                    if st.button("💾 保存配置", type="primary", use_container_width=True):  
                        save_sops()  
                        st.success("已保存！")  
  
                if new_name != s_name and new_name.strip():  
                    st.session_state.sops[new_name] = st.session_state.sops.pop(s_name)  
                    if s_name in st.session_state.memory:  
                        st.session_state.memory[new_name] = st.session_state.memory.pop(s_name)  
                        save_memory()  
                    save_sops()  
                    st.rerun()  
  
                st.markdown("### 🧠 记忆生长模式")  
                mode_opts = {  
                    "manual": "保守派：手动挑选好文章提取风格",  
                    "dynamic": "激进派：自动活体记忆避坑反馈"  
                }  
                current_mode_idx = 0 if sop.get("memory_mode", "manual") == "manual" else 1  
                sop["memory_mode"] = st.radio(  
                    "选择成长路线", ["manual", "dynamic"],  
                    format_func=lambda x: mode_opts[x], index=current_mode_idx,  
                    label_visibility="collapsed"  
                )  
  
                if sop["memory_mode"] == "dynamic" and sop.get("negative_memory"):  
                    with st.expander("👀 已吸收的避坑清单", expanded=True):  
                        for nm in sop["negative_memory"]:  
                            st.markdown("- " + nm)  
                        if st.button("强行清空避坑清单"):  
                            sop["negative_memory"] = []  
                            save_sops()  
                            st.rerun()  
  
                st.markdown("### 🎭 专属人设 (System Prompt)")  
                sop["system_prompt"] = st.text_area("核心指令", sop.get("system_prompt", ""), height=100)  
  
                st.markdown("### 🧩 执行阶段配置")  
                new_steps = []  
                for i, step in enumerate(sop["steps"]):  
                    _ensure_step(step)  
                    with st.container(border=True):  
                        st.markdown("**阶段 " + str(i + 1) + "**")  
                        sc1, sc2 = st.columns([4, 1])  
                        with sc1:  
                            p_val = st.text_area(  
                                "指令", step["prompt"], height=60,  
                                key="p_" + str(i), label_visibility="collapsed"  
                            )  
                        with sc2:  
                            l_val = st.number_input(  
                                "循环", min_value=1, value=step.get("loop", 1), key="l_" + str(i)  
                            )  
  
                        # 字数精控（可选）  
                        wc_on = st.checkbox(  
                            "📏 启用字数精控", step.get("enable_word_control", False),  
                            key="wc_on_" + str(i)  
                        )  
                        tgt_w = step.get("target_words", 0)  
                        tol_w = step.get("word_tolerance", 5)  
                        mc_w = step.get("max_corrections", 2)  
                        if wc_on:  
                            wcc1, wcc2, wcc3 = st.columns(3)  
                            with wcc1:  
                                tgt_w = st.number_input(  
                                    "目标字数", min_value=100, value=max(tgt_w, 100),  
                                    step=100, key="tgt_" + str(i)  
                                )  
                            with wcc2:  
                                tol_w = st.slider(  
                                    "容差%", 1, 30, tol_w, key="tol_" + str(i)  
                                )  
                            with wcc3:  
                                mc_w = st.number_input(  
                                    "最大矫正轮", min_value=0, max_value=5,  
                                    value=mc_w, key="mc_" + str(i)  
                                )  
  
                        with st.expander("📁 挂载阶段参考资料"):  
                            ref = st.text_area(  
                                "粘贴资料", step.get("reference", ""), key="r_" + str(i)  
                            )  
  
                        new_steps.append({  
                            "prompt": p_val, "loop": l_val, "reference": ref,  
                            "enable_word_control": wc_on, "target_words": tgt_w,  
                            "word_tolerance": tol_w, "max_corrections": mc_w  
                        })  
                sop["steps"] = new_steps  
  
                sbc1, sbc2 = st.columns(2)  
                with sbc1:  
                    if st.button("➕ 加阶段", use_container_width=True):  
                        sop["steps"].append(_default_step())  
                        save_sops()  
                        st.rerun()  
                with sbc2:  
                    if len(sop["steps"]) > 1 and st.button("➖ 删末尾阶段", use_container_width=True):  
                        sop["steps"].pop()  
                        save_sops()  
                        st.rerun()  
  
                st.markdown("### ⚡ 监听触发器网络")  
                new_triggers = []  
                for i, t in enumerate(sop.get("triggers", [])):  
                    with st.container(border=True):  
                        tc1, tc2, tc3 = st.columns([1, 1, 2])  
                        with tc1:  
                            typ = st.selectbox(  
                                "类型", ["terminate", "intervene"],  
                                index=0 if t["type"] == "terminate" else 1,  
                                key="t_" + str(i)  
                            )  
                        with tc2:  
                            kwd = st.text_input("关键词", t["keyword"], key="k_" + str(i))  
                        with tc3:  
                            act = st.text_input(  
                                "动作指令", t.get("action", ""),  
                                disabled=(typ == "terminate"), key="a_" + str(i)  
                            )  
                        new_triggers.append({"type": typ, "keyword": kwd, "action": act})  
                sop["triggers"] = new_triggers  
                if st.button("➕ 加规则"):  
                    sop["triggers"].append({"type": "intervene", "keyword": "", "action": ""})  
                    save_sops()  
                    st.rerun()  
  
                st.divider()  
                # 删除 SOP（二次确认）  
                if st.session_state.get("_confirm_del_sop"):  
                    st.error("确认删除 SOP [" + s_name + "]？不可撤回！")  
                    dsc1, dsc2 = st.columns(2)  
                    if dsc1.button("✅ 确认删除", key="yes_del_sop"):  
                        del st.session_state.sops[s_name]  
                        if s_name in st.session_state.memory:  
                            del st.session_state.memory[s_name]  
                            save_memory()  
                        save_sops()  
                        st.session_state._confirm_del_sop = False  
                        st.rerun()  
                    if dsc2.button("❌ 取消", key="no_del_sop"):  
                        st.session_state._confirm_del_sop = False  
                        st.rerun()  
                else:  
                    if st.button("🗑️ 删除此 SOP", type="primary"):  
                        st.session_state._confirm_del_sop = True  
                        st.rerun()  
  
    with tab_vault:  
        st.header("🗄️ 记忆保险库与炼丹炉")  
        if not sop_list:  
            st.info("请先创建 SOP")  
        else:  
            vault_sop = st.selectbox("选择账号", sop_list, key="vault_sel")  
            acc_mem = st.session_state.memory.get(vault_sop, [])  
            if not acc_mem:  
                st.info("【" + vault_sop + "】保险库为空。")  
            else:  
                st.success("🗃️ 已沉淀 " + str(len(acc_mem)) + " 篇作品。")  
                for idx, item in enumerate(reversed(acc_mem)):  
                    with st.expander("📖 [" + item['time'] + "] " + item['topic']):  
                        st.caption(item['content'])  
                        if st.button("🗑️ 抹除此记忆", key="del_mem_" + str(idx)):  
                            st.session_state.memory[vault_sop].remove(item)  
                            save_memory()  
                            st.rerun()  
  
                st.divider()  
                if st.button("🔥 立即开炉提炼灵魂 (风格蒸馏)", type="primary", use_container_width=True):  
                    if not active_p["api_key"]:  
                        st.error("缺 API Key！")  
                    else:  
                        client, profile = get_client()  
                        combined = "\n\n---\n\n".join([m['content'] for m in acc_mem[-3:]])  
                        distill_p = (  
                            "深度分析以下小说风格，提炼一段极度严谨的【System Prompt】复刻文风。"  
                            "只需输出纯文本的 Prompt。\n样本：\n" + combined  
                        )  
                        with st.spinner("正在提炼灵魂设定..."):  
                            try:  
                                resp = client.chat.completions.create(  
                                    model=profile["model"],  
                                    messages=[{"role": "user", "content": distill_p}]  
                                )  
                                result = resp.choices[0].message.content.strip()  
                                st.session_state.sops[vault_sop]["system_prompt"] = result  
                                save_sops()  
                                st.success("🎉 蒸馏成功！人设已注入全新灵魂！")  
                                st.info("**提炼成果:**\n\n" + result)  
                            except Exception as e:  
                                st.error("蒸馏失败: " + str(e))  
  
# ==========================================  
# 模块 4: 底层引擎配置  
# ==========================================  
elif st.session_state.current_page == "⚙️ 底层引擎配置":  
    st.header("⚙️ 底层驱动配置")  
    col_list, col_edit = st.columns([1, 2.5])  
  
    with col_list:  
        st.subheader("引擎库")  
        p_names = [p["name"] for p in st.session_state.profiles]  
        idx = st.radio(  
            "切换引擎", range(len(p_names)),  
            format_func=lambda x: p_names[x],  
            index=st.session_state.active_profile_idx  
        )  
        st.session_state.active_profile_idx = idx  
        if st.button("➕ 新增引擎", use_container_width=True):  
            st.session_state.profiles.append({  
                "name": "新引擎 " + str(len(p_names) + 1),  
                "base_url": "", "api_key": "",  
                "model": "anthropic/claude-sonnet-4.6",  
                "use_temperature": True, "temperature": 0.8,  
                "use_max_tokens": True, "max_tokens": 4096,  
                "use_top_p": False, "top_p": 1.0,  
                "use_frequency_penalty": False, "frequency_penalty": 0.0  
            })  
            save_profiles()  
            st.rerun()  
  
    with col_edit:  
        st.subheader("网络与参数调优")  
        p = st.session_state.profiles[idx]  
  
        eca, ecb = st.columns([3, 1])  
        with eca:  
            p["name"] = st.text_input("引擎标签", p["name"])  
        with ecb:  
            st.write("")  
            if st.button("💾 保存引擎配置", type="primary", use_container_width=True):  
                save_profiles()  
                st.success("已保存！")  
  
        nc1, nc2 = st.columns(2)  
        with nc1:  
            p["base_url"] = st.text_input("Base URL", p["base_url"])  
        with nc2:  
            p["api_key"] = st.text_input("API Key", p["api_key"], type="password")  
  
        # API Key 连通测试  
        if p["api_key"]:  
            if st.button("🔑 测试连通性"):  
                with st.spinner("正在测试..."):  
                    try:  
                        test_client = OpenAI(  
                            base_url=p["base_url"].strip() or "https://api.openai.com/v1",  
                            api_key=p["api_key"].strip()  
                        )  
                        test_resp = test_client.chat.completions.create(  
                            model=p["model"],  
                            messages=[{"role": "user", "content": "Hi"}],  
                            max_tokens=5  
                        )  
                        st.success("✅ 连通成功！模型响应正常。")  
                    except Exception as e:  
                        st.error("❌ 连通失败: " + str(e))  
  
        cm, cb2 = st.columns([3, 1])  
        with cm:  
            p["model"] = st.text_input("模型映射 (Model ID)", p["model"])  
        with cb2:  
            st.write("")  
            if st.button("🔄 联机获取列表"):  
                if p["api_key"]:  
                    with st.spinner("正在获取..."):  
                        success, result = fetch_models(p["base_url"], p["api_key"])  
                        if success:  
                            if result:  
                                st.session_state.temp_models = result  
                                st.success("✅ 抓取到 " + str(len(result)) + " 个模型！")  
                            else:  
                                st.warning("请求成功但无模型列表。")  
                        else:  
                            st.error("❌ 获取失败: " + str(result))  
                else:  
                    st.error("请先填写 API Key！")  
  
        if "temp_models" in st.session_state:  
            sel_m = st.selectbox("选择模型", ["(不覆盖)"] + st.session_state.temp_models)  
            if sel_m != "(不覆盖)":  
                p["model"] = sel_m  
                del st.session_state.temp_models  
                save_profiles()  
                st.rerun()  
  
        st.markdown("#### 🎛️ 运行时超参数 (勾选生效)")  
        sl1, sl2 = st.columns(2)  
        with sl1:  
            p["use_temperature"] = st.checkbox("🔥 Temperature", p.get("use_temperature", True))  
            if p["use_temperature"]:  
                p["temperature"] = st.slider("温度值", 0.0, 2.0, p.get("temperature", 0.8), 0.1, label_visibility="collapsed")  
            p["use_max_tokens"] = st.checkbox("📏 Max Tokens", p.get("use_max_tokens", True))  
            if p["use_max_tokens"]:  
                p["max_tokens"] = st.slider("最大Token", 512, 16384, p.get("max_tokens", 4096), 512, label_visibility="collapsed")  
        with sl2:  
            p["use_top_p"] = st.checkbox("🎲 Top P", p.get("use_top_p", False))  
            if p["use_top_p"]:  
                p["top_p"] = st.slider("Top P值", 0.0, 1.0, p.get("top_p", 1.0), 0.05, label_visibility="collapsed")  
            p["use_frequency_penalty"] = st.checkbox("🚫 Frequency Penalty", p.get("use_frequency_penalty", False))  
            if p["use_frequency_penalty"]:  
                p["frequency_penalty"] = st.slider("惩罚值", -2.0, 2.0, p.get("frequency_penalty", 0.0), 0.1, label_visibility="collapsed")  
  
        # 删除引擎（二次确认）  
        if len(st.session_state.profiles) > 1:  
            st.divider()  
            if st.session_state.get("_confirm_del_engine"):  
                st.error("确认删除引擎 [" + p['name'] + "]？")  
                dec1, dec2 = st.columns(2)  
                if dec1.button("✅ 确认", key="yes_del_eng"):  
                    st.session_state.profiles.pop(idx)  
                    st.session_state.active_profile_idx = 0  
                    st.session_state._confirm_del_engine = False  
                    save_profiles()  
                    st.rerun()  
                if dec2.button("❌ 取消", key="no_del_eng"):  
                    st.session_state._confirm_del_engine = False  
                    st.rerun()  
            else:  
                if st.button("🗑️ 删除此引擎"):  
                    st.session_state._confirm_del_engine = True  
                    st.rerun()  
