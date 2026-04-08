import streamlit as st
from openai import OpenAI
import io
import os
import json
import uuid
from datetime import datetime
from docx import Document

# ==========================================
# 1. 页面全局配置与初始化
# ==========================================
st.set_page_config(page_title="ZenMux AI 助手", page_icon="🦀", layout="wide")

HISTORY_DIR = "history"
CONFIG_FILE = "config.json"

# 确保历史记录文件夹存在
if not os.path.exists(HISTORY_DIR):
    os.makedirs(HISTORY_DIR)

# ==========================================
# 2. 核心功能：数据持久化与文件处理
# ==========================================
def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"api_key": ""}

def save_config(api_key):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump({"api_key": api_key}, f)

def get_all_sessions():
    sessions = []
    for filename in os.listdir(HISTORY_DIR):
        if filename.endswith(".json"):
            filepath = os.path.join(HISTORY_DIR, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                sessions.append(data)
    # 按时间倒序排列
    return sorted(sessions, key=lambda x: x['updated_at'], reverse=True)

def save_session(session_id, title, messages):
    filepath = os.path.join(HISTORY_DIR, f"{session_id}.json")
    data = {
        "id": session_id,
        "title": title,
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "messages": messages
    }
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def delete_session(session_id):
    filepath = os.path.join(HISTORY_DIR, f"{session_id}.json")
    if os.path.exists(filepath):
        os.remove(filepath)

def generate_word_doc(messages):
    doc = Document()
    doc.add_heading('ZenMux AI 对话记录', 0)
    for msg in messages:
        if msg["role"] == "system": continue
        role_name = "🧑‍💻 我" if msg["role"] == "user" else "🤖 AI"
        doc.add_heading(role_name, level=2)
        doc.add_paragraph(msg["content"])
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 3. 状态管理 (Session State)
# ==========================================
if "config" not in st.session_state:
    st.session_state.config = load_config()

if "current_session_id" not in st.session_state:
    st.session_state.current_session_id = None

if "messages" not in st.session_state:
    st.session_state.messages = []

# 初始化加载历史会话列表
all_sessions = get_all_sessions()

# 如果没有当前会话，且有历史会话，加载第一个
if st.session_state.current_session_id is None and all_sessions:
    st.session_state.current_session_id = all_sessions[0]["id"]
    st.session_state.messages = all_sessions[0]["messages"]
# 如果没有任何会话，新建一个
elif st.session_state.current_session_id is None:
    st.session_state.current_session_id = str(uuid.uuid4())

# ==========================================
# 4. 侧边栏：历史记录、配置与导出
# ==========================================
with st.sidebar:
    st.header("💬 会话管理")
    if st.button("➕ 新建对话", use_container_width=True):
        st.session_state.current_session_id = str(uuid.uuid4())
        st.session_state.messages = []
        st.rerun()

    st.subheader("历史记录")
    for sess in all_sessions:
        col1, col2 = st.columns([4, 1])
        with col1:
            # 截断过长的标题
            display_title = sess['title'][:12] + "..." if len(sess['title']) > 12 else sess['title']
            if st.button(f"{display_title}", key=f"btn_{sess['id']}", use_container_width=True):
                st.session_state.current_session_id = sess['id']
                st.session_state.messages = sess['messages']
                st.rerun()
        with col2:
            if st.button("🗑", key=f"del_{sess['id']}"):
                delete_session(sess['id'])
                if st.session_state.current_session_id == sess['id']:
                    st.session_state.current_session_id = None
                    st.session_state.messages = []
                st.rerun()

    st.divider()
    
    st.header("⚙️ 核心配置")
    api_key = st.text_input(
        "🔑 ZenMux API Key", 
        type="password", 
        value=st.session_state.config.get("api_key", ""),
        placeholder="sk-ai-v1-xxxxxxxxxxxx...",
    )
    if api_key != st.session_state.config.get("api_key", ""):
        save_config(api_key)
        st.session_state.config["api_key"] = api_key

    model_name = st.selectbox(
        "🧠 选择模型", 
        ["anthropic/claude-opus-4.6", "anthropic/claude-sonnet-4.6", "anthropic/claude-haiku-4.5"]
    )
    
    st.divider()
    
    st.header("📄 附加资料区")
    uploaded_file = st.file_uploader("上传参考文档 (txt, md, csv)", type=['txt', 'md', 'csv'])
    file_content = ""
    if uploaded_file is not None:
        try:
            file_content = uploaded_file.getvalue().decode("utf-8")
            st.success(f"✅ [{uploaded_file.name}] 已作为背景知识载入")
        except Exception as e:
            st.error(f"❌ 解析失败: {e}")

    st.divider()
    
    st.header("📦 导出当前对话")
    if not st.session_state.messages:
        st.info("暂无对话可导出")
    else:
        # 生成 TXT 字符串
        txt_content = ""
        for m in st.session_state.messages:
            if m["role"] == "system": continue
            role = "我" if m["role"] == "user" else "AI"
            txt_content += f"{role}:\n{m['content']}\n\n{'-'*40}\n\n"
            
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("📥 TXT格式", txt_content, "AI对话记录.txt", "text/plain")
        with c2:
            st.download_button("📥 Word格式", generate_word_doc(st.session_state.messages), "AI对话记录.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==========================================
# 5. 主界面：聊天窗口
# ==========================================
st.title("🤖 ZenMux 多模态智能助手")

if not api_key:
    st.warning("👈 请先在左侧边栏输入您的 API Key 以激活对话功能。")
    st.stop()

try:
    client = OpenAI(base_url="https://zenmux.ai/api/v1", api_key=api_key)
except Exception as e:
    st.error(f"客户端初始化失败: {e}")
    st.stop()

# 渲染当前对话
for message in st.session_state.messages:
    if message["role"] == "system": continue
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# 处理用户输入
if prompt := st.chat_input(f"想问点什么？(当前使用: {model_name})"):
    # 显示用户输入
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
        
    # 构建 API 请求
    api_messages = []
    if file_content:
        system_msg = f"你是一个强大的 AI 助手。请基于以下用户上传的文件内容来回答问题。文件内容：\n\n{file_content}"
        api_messages.append({"role": "system", "content": system_msg})
        
    api_messages.extend([{"role": m["role"], "content": m["content"]} for m in st.session_state.messages])

    with st.chat_message("assistant"):
        try:
            stream = client.chat.completions.create(
                model=model_name,
                messages=api_messages,
                stream=True
            )
            response = st.write_stream(stream)
            st.session_state.messages.append({"role": "assistant", "content": response})
            
            # 生成对话标题（取第一句话的前15个字）
            title = st.session_state.messages[0]["content"][:15] if len(st.session_state.messages) > 0 else "新对话"
            # 保存到本地历史记录
            save_session(st.session_state.current_session_id, title, st.session_state.messages)
            
            st.rerun()
        except Exception as e:
            st.error(f"请求失败，请检查配置。\n错误详情: {e}")
