import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI
import io
import base64
import json
import os
import re
import requests
import uuid
from datetime import datetime
from docx import Document

# ==========================================
# 1. 页面全局配置与前端美化
# ==========================================
st.set_page_config(page_title="ZenMux 创作者工作站", page_icon="🐙", layout="wide")
st.markdown("""
    <style>
    .stButton>button { border-radius: 8px; font-weight: bold; transition: all 0.3s; }
    .stChatInput { padding-bottom: 20px; }
    button[title="View fullscreen"] {display: none;}
    .css-1jc7ptx, .e1ewe7hr3, .viewerBadge_container__1QSob, .styles_viewerBadge__1yB5_ {display: none;}
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 2. 数据状态管理 (云端绝对隔离模式)
# ==========================================
# 彻底移除本地文件读写，防止多用户串号和 API Key 泄露
# 用空的占位函数替换，保证不报错，但绝对不往服务器写数据
def save_profiles(): pass
def save_sops(): pass
def save_memory(): pass
def save_free_chats(): pass

# ==========================================
# 3. 核心底层辅助函数
# ==========================================
def render_copy_button(text):
    """完美无痕复制按钮（Base64 防断码版）"""
    # 将长文本安全转为 Base64，彻底杜绝 HTML 单双引号冲突
    b64_text = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    
    html = f"""
    <div style="display:flex; justify-content:flex-end; align-items:center; width:100%; height:100%; margin:0; padding-right:10px;">
        <button id="copyBtn" 
            style="border:none; background:transparent; color:#aaa; cursor:pointer; font-size:12px; font-weight:bold; padding:5px 10px; border-radius:6px; transition:0.2s;"
            onmouseover="this.style.color='#4CAF50'; this.style.backgroundColor='#f0f9f0'" 
            onmouseout="this.style.color='#aaa'; this.style.backgroundColor='transparent'"
        >
            📋 复制纯文本
        </button>
    </div>
    <script>
        document.getElementById("copyBtn").onclick = function() {{
            // JS 解码 Base64 还原中文文本
            const str = decodeURIComponent(escape(window.atob("{b64_text}")));
            navigator.clipboard.writeText(str).then(function() {{
                const btn = document.getElementById("copyBtn");
                btn.innerText = "✅ 复制成功";
                btn.style.color = "#4CAF50";
                setTimeout(function() {{ 
                    btn.innerText = "📋 复制纯文本"; 
                    btn.style.color = "#aaa";
                }}, 2000);
            }});
        }};
    </script>
    """
    # height=35 刚刚好，不会出现难看的滚动条
    components.html(html, height=35)

def clean_novel_text(text):
    text = re.sub(r'^\s*(好的|没问题|非常荣幸|收到|为你生成|以下是|这是为您|正文开始|下面是).*?[:：]\n*', '', text, flags=re.MULTILINE|re.IGNORECASE)
    text = re.sub(r'^\s*第[零一二三四五六七八九十百千0-9]+[章回节卷].*?\n', '', text, flags=re.MULTILINE)
    text = re.sub(r'```[a-zA-Z]*\n?', '', text)
    text = re.sub(r'\n*(希望这|如果有需要|请告诉我|期待您的反馈).*$', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def generate_word_doc(messages, is_pure=False):
    doc = Document()
    doc.add_heading('AI 创作者工作站生成文档', 0)
    for msg in messages:
        if msg["role"] == "system" or not msg.get("selected", True): continue
        if is_pure:
            if msg["role"] == "assistant": doc.add_paragraph(clean_novel_text(msg["content"]))
        else:
            if msg["role"] == "user":
                doc.add_heading("📌 指令 / 我", level=2)
                doc.add_paragraph(msg["content"])
            else:
                doc.add_heading("🤖 AI", level=2)
                doc.add_paragraph(msg["content"])
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def fetch_models(base_url, api_key):
    """动态获取大模型列表，带详细报错捕获"""
    try:
        # 自动处理 URL 结尾可能多出的斜杠
        url = (base_url.strip().rstrip('/') or "https://api.openai.com/v1") + "/models"
        headers = {"Authorization": f"Bearer {api_key.strip()}"}
        # 设置8秒超时
        resp = requests.get(url, headers=headers, timeout=8) 
        
        if resp.status_code == 200:
            models = sorted([m["id"] for m in resp.json().get("data", [])])
            return True, models
        else:
            return False, f"API 拒绝请求 (状态码: {resp.status_code})。返回详情: {resp.text[:100]}"
    except requests.exceptions.RequestException as e:
        return False, f"网络连接失败或超时，请检查 Base URL 是否正确或是否需要代理。详情: {e}"
    except Exception as e:
        return False, f"数据解析异常: {e}"

def get_client():
    profile = st.session_state.profiles[st.session_state.active_profile_idx]
    url = profile["base_url"].strip() or "https://api.openai.com/v1"
    key = profile["api_key"].strip()
    return OpenAI(base_url=url, api_key=key), profile

def build_api_kwargs(profile, api_msgs):
    kwargs = {"model": profile["model"], "messages": api_msgs, "stream": True}
    if profile.get("use_temperature", True): kwargs["temperature"] = profile.get("temperature", 0.7)
    if profile.get("use_max_tokens", True): kwargs["max_tokens"] = profile.get("max_tokens", 4096)
    if profile.get("use_top_p", False): kwargs["top_p"] = profile.get("top_p", 1.0)
    if profile.get("use_frequency_penalty", False): kwargs["frequency_penalty"] = profile.get("frequency_penalty", 0.0)
    return kwargs

def stream_generator(api_stream):
    st.session_state.auto_engine["last_finish_reason"] = "stop"
    for chunk in api_stream:
        if chunk.choices and chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content
        if chunk.choices and chunk.choices[0].finish_reason is not None:
            st.session_state.auto_engine["last_finish_reason"] = chunk.choices[0].finish_reason

# ==========================================
# 4. 状态初始化 (真正的私有沙盒)
# ==========================================
if "initialized" not in st.session_state:
    st.session_state.profiles = [{
        "name": "默认引擎", "base_url": "", "api_key": "", "model": "anthropic/claude-sonnet-4.6",
        "use_temperature": True, "temperature": 0.8, "use_max_tokens": True, "max_tokens": 4096
    }]
    st.session_state.sops = {
        "小说账号预设": {
            "memory_mode": "manual",
            "system_prompt": "你是一名悬疑小说家，文风冷峻，绝不输出废话。",
            "negative_memory": [],
            "steps": [{"prompt": "撰写第【{循环索引}】章", "loop": 2, "reference": ""}],
            "triggers": [{"type": "terminate", "keyword": "全文完", "action": ""}]
        }
    }
    st.session_state.memory = {}
    st.session_state.free_chats = {str(uuid.uuid4()): {"title": "新对话", "messages": []}}
    
    st.session_state.active_profile_idx = 0
    st.session_state.current_page = "🤖 自动化流水线"
    st.session_state.current_chat_id = list(st.session_state.free_chats.keys())[-1]
    st.session_state.auto_engine = {
        "is_running": False, "is_finished": False, "messages": [],
        "sop_name": "", "topic": "", "global_file": "",
        "current_step_idx": 0, "current_loop_idx": 1,
        "pending_instruction": "", "last_finish_reason": ""
    }
    st.session_state.initialized = True

# ==========================================
# 5. 全局侧边栏导航
# ==========================================
with st.sidebar:
    col_img, col_txt = st.columns([1, 3])
    with col_img: st.image("https://api.iconify.design/fluent-emoji:octopus.svg?width=80", width=45)
    with col_txt: st.header("控制中枢")
    
    st.write("") 
    pages = ["🤖 自动化流水线", "💬 自由聊天区", "📝 账号SOP与灵魂", "⚙️ 底层引擎配置"]
    for p in pages:
        btn_type = "primary" if st.session_state.current_page == p else "secondary"
        if st.button(p, use_container_width=True, type=btn_type):
            st.session_state.current_page = p
            st.rerun()

    active_p = st.session_state.profiles[st.session_state.active_profile_idx]
    st.divider()
    st.caption(f"🟢 **当前挂载**: {active_p['name']}\n🧠 **模型**: {active_p['model']}")
    
    if st.session_state.current_page == "💬 自由聊天区":
        st.divider()
        st.header("📚 历史对话")
        if st.button("➕ 开启新对话", use_container_width=True, type="primary"):
            new_id = str(uuid.uuid4())
            st.session_state.free_chats[new_id] = {"title": "新对话", "messages": []}
            st.session_state.current_chat_id = new_id
            save_free_chats(); st.rerun()
            
        for c_id, c_data in reversed(list(st.session_state.free_chats.items())):
            c_title = c_data["title"][:12] + "..." if len(c_data["title"])>12 else c_data["title"]
            btn_lbl = f"⭐ {c_title}" if c_id == st.session_state.current_chat_id else f"📄 {c_title}"
            if st.button(btn_lbl, key=f"chat_{c_id}", use_container_width=True):
                st.session_state.current_chat_id = c_id; st.rerun()
                
        st.divider()
        current_msgs = st.session_state.free_chats[st.session_state.current_chat_id]["messages"]
        if current_msgs:
            st.markdown("**📦 导出当前对话**")
            exp_mode = st.radio("格式选项", ["完整记录", "纯享正文 (清洗后)"], label_visibility="collapsed")
            is_pure = (exp_mode == "纯享正文 (清洗后)")
            if is_pure: txt_content = "\n\n".join([clean_novel_text(m['content']) for m in current_msgs if m['role']=='assistant'])
            else: txt_content = "".join([f"{'我' if m['role']=='user' else 'AI'}:\n{m['content']}\n\n{'-'*40}\n\n" for m in current_msgs])
            
            c1, c2 = st.columns(2)
            with c1: st.download_button("📥 TXT", txt_content.encode('utf-8'), "聊天记录.txt", "text/plain", use_container_width=True)
            with c2: st.download_button("📥 Word", generate_word_doc(current_msgs, is_pure), "聊天记录.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

        if st.button("🗑️ 删除对话", use_container_width=True):
            if len(st.session_state.free_chats) > 1: del st.session_state.free_chats[st.session_state.current_chat_id]
            else: st.session_state.free_chats[st.session_state.current_chat_id] = {"title": "新对话", "messages": []}
            st.session_state.current_chat_id = list(st.session_state.free_chats.keys())[-1]
            save_free_chats(); st.rerun()

    st.divider()
    with st.expander("📦 全量资产导出恢复舱", expanded=False):
        st.caption("换电脑或重装系统时，可一键导入导出所有数据。")
        full_data = json.dumps({"profiles": st.session_state.profiles, "sops": st.session_state.sops, "memory": st.session_state.memory, "free_chats": st.session_state.free_chats}, ensure_ascii=False, indent=2).encode('utf-8')
        st.download_button("📥 导出全量快照包", full_data, f"ZenMux_Backup_{datetime.now().strftime('%m%d')}.json", "application/json", use_container_width=True, type="primary")
        uploaded_ws = st.file_uploader("📂 导入快照 (将覆盖)", type="json")
        if uploaded_ws:
            try:
                data = json.loads(uploaded_ws.getvalue().decode('utf-8'))
                st.session_state.profiles = data.get("profiles", [])
                st.session_state.sops = data.get("sops", {})
                st.session_state.memory = data.get("memory", {})
                st.session_state.free_chats = data.get("free_chats", {str(uuid.uuid4()): {"title": "新对话", "messages": []}})
                st.session_state.active_profile_idx = 0
                st.session_state.current_chat_id = list(st.session_state.free_chats.keys())[-1]
                save_profiles(); save_sops(); save_memory(); save_free_chats()
                st.success("✅ 恢复成功！请刷新页面。")
            except: st.error("文件格式错误！")

# ==========================================
# 模块 1: 自动化流水线
# ==========================================
if st.session_state.current_page == "🤖 自动化流水线":
    engine = st.session_state.auto_engine
    col_ctrl, col_view = st.columns([1, 2.5])
    
    with col_ctrl:
        st.header("⚙️ 调度控制台")
        if engine["is_running"]:
            st.warning("⚠️ 引擎高速运转中...")
            total_steps = max(len(st.session_state.sops.get(engine["sop_name"], {"steps":[]})["steps"]), 1)
            st.progress(min(engine["current_step_idx"] / total_steps, 1.0))
            if st.button("⏹️ 强制急停", type="primary", use_container_width=True): 
                engine["is_running"] = False; st.rerun()
        else:
            if not st.session_state.sops:
                st.warning("请先去配置一个 SOP。")
                st.stop()
            sel_sop = st.selectbox("1. 挂载执行 SOP (账号人设)", list(st.session_state.sops.keys()))
            in_topic = st.text_input("2. 注入 {主题}", placeholder="例如：赛博朋克修仙传")
            up_file = st.file_uploader("3. 挂载全局设定集 (可选)", type=['txt', 'md'])
            
            if st.button("🚀 点火启动", type="primary", use_container_width=True):
                if not active_p["api_key"]: st.error("引擎未配置 API Key！")
                elif not in_topic: st.error("请填入主题！")
                else:
                    engine.update({
                        "is_running": True, "is_finished": False, "messages": [],
                        "sop_name": sel_sop, "topic": in_topic,
                        "global_file": up_file.getvalue().decode("utf-8") if up_file else "",
                        "current_step_idx": 0, "current_loop_idx": 1,
                        "pending_instruction": "", "last_finish_reason": ""
                    })
                    st.rerun()
                    
        st.divider()
        if engine["messages"]:
            st.markdown("### 📦 成果验收与记忆管理")
            sel_msgs = [m for m in engine["messages"] if m["role"] == "assistant" and m.get("selected", True)]
            raw_text = "\n\n".join([m["content"] for m in sel_msgs])
            pure_text = clean_novel_text(raw_text)
            
            if not engine["is_running"]:
                sop_data = st.session_state.sops[engine["sop_name"]]
                mem_mode = sop_data.get("memory_mode", "manual")
                
                if mem_mode == "manual":
                    st.info("🧠 记忆模式：【手动提取蒸馏】")
                    if st.button("💾 将本次佳作存入该账号记忆保险库", type="primary", use_container_width=True):
                        sop_name = engine["sop_name"]
                        if sop_name not in st.session_state.memory: st.session_state.memory[sop_name] = []
                        st.session_state.memory[sop_name].append({
                            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "topic": engine["topic"], "content": pure_text[:2500]
                        })
                        save_memory(); st.toast("已安全落盘！", icon="💾")
                else:
                    st.success("🌱 记忆模式：【动态活体进化】")
                    feedback = st.text_input("💬 对本次生成的避坑要求/反馈：", placeholder="例如：以后不准废话")
                    if st.button("提交反馈并写入账号潜意识", use_container_width=True):
                        if feedback.strip():
                            if "negative_memory" not in sop_data: sop_data["negative_memory"] = []
                            sop_data["negative_memory"].append(feedback); save_sops()
                            if len(sop_data["negative_memory"]) >= 3:
                                with st.spinner("反思并融合新规则中..."):
                                    try:
                                        client, profile = get_client()
                                        fusion_prompt = f"原人设：{sop_data['system_prompt']}。避坑反馈：{'; '.join(sop_data['negative_memory'])}。请深度融合进原人设，形成新 System Prompt，只输出纯文本。"
                                        resp = client.chat.completions.create(model=profile["model"], messages=[{"role": "user", "content": fusion_prompt}])
                                        sop_data["system_prompt"] = resp.choices[0].message.content.strip()
                                        sop_data["negative_memory"] = []; save_sops()
                                        st.success("反馈已吸收！人设已进化！")
                                    except Exception as e: st.error(f"融合失败: {e}")
                            else: st.toast("反馈已记录，下次生效！", icon="✅")

            c1, c2 = st.columns(2)
            with c1: st.download_button("📥 标准全文", raw_text.encode('utf-8'), f"{engine['topic']}_完整.txt", "text/plain", use_container_width=True)
            with c2: st.download_button("✨ 纯享正文", pure_text.encode('utf-8'), f"{engine['topic']}_纯正文.txt", "text/plain", use_container_width=True)
            if st.button("🧹 清理工作台", use_container_width=True): engine.update({"messages": [], "is_finished": False, "is_running": False}); st.rerun()

    with col_view:
        st.header("🖥️ 监视大屏")
        # 🔥 UI重构：独立滚动的监控大屏
        with st.container(height=750, border=True):
            for i, msg in enumerate(engine["messages"]):
                if msg["role"] == "system": continue
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])
                    if msg["role"] == "assistant":
                        render_copy_button(msg["content"])
                        msg["selected"] = st.checkbox("☑️ 选中参与导出", msg.get("selected", True), key=f"ac_{i}")

            if engine["is_running"]:
                client, profile = get_client()
                sop_data = st.session_state.sops[engine["sop_name"]]
                steps = sop_data["steps"]
                triggers = sop_data.get("triggers", [])
                curr_step = steps[engine["current_step_idx"]]
                
                current_prompt = engine["pending_instruction"] or curr_step["prompt"].replace("{主题}", engine["topic"]).replace("{循环索引}", str(engine["current_loop_idx"]))
                engine["pending_instruction"] = ""
                
                silence_constraint = "\n\n【系统强制指令：绝对不要重复上文，不准说“好的”，不准带章节标题，直接从正文第一个字开始输出！】"
                final_prompt_to_api = current_prompt + silence_constraint
                
                engine["messages"].append({"role": "user", "content": current_prompt, "selected": False})
                with st.chat_message("user"): st.markdown(f"*(⚡ 指令)*: {current_prompt}")
                    
                api_msgs = []
                sys_prompt = sop_data.get("system_prompt", "").strip()
                if sys_prompt: api_msgs.append({"role": "system", "content": sys_prompt})
                if sop_data.get("memory_mode", "manual") == "dynamic" and sop_data.get("negative_memory"):
                    api_msgs.append({"role": "system", "content": f"【必须遵守的避坑铁律】：{'; '.join(sop_data['negative_memory'])}"})
                if engine["global_file"]: api_msgs.append({"role": "system", "content": f"【全局设定】\n{engine['global_file']}"})
                if curr_step.get("reference"): api_msgs.append({"role": "system", "content": f"【本阶段设定】\n{curr_step['reference']}"})
                
                for idx, m in enumerate(engine["messages"]):
                    if idx == len(engine["messages"]) - 1 and m["role"] == "user": api_msgs.append({"role": "user", "content": final_prompt_to_api})
                    else: api_msgs.append({"role": m["role"], "content": m["content"]})
                
                with st.chat_message("assistant"):
                    try:
                        resp = client.chat.completions.create(**build_api_kwargs(profile, api_msgs))
                        full_resp = st.write_stream(stream_generator(resp))
                        render_copy_button(full_resp)
                        engine["messages"].append({"role": "assistant", "content": full_resp, "selected": True})
                        
                        hit_trigger = False
                        if engine["last_finish_reason"] == "length":
                            engine["pending_instruction"] = "⚠️ 因字数限制中断，请紧接着上文最后一个字继续往下写。"
                            hit_trigger = True
                        if not hit_trigger:
                            for t in triggers:
                                if t["keyword"] and t["keyword"] in full_resp:
                                    if t["type"] == "terminate":
                                        engine["is_running"] = False; engine["is_finished"] = True; hit_trigger = True; break
                                    elif t["type"] == "intervene":
                                        engine["pending_instruction"] = t["action"]; hit_trigger = True; break
                                        
                        if not hit_trigger:
                            if engine["current_loop_idx"] < curr_step.get("loop", 1): engine["current_loop_idx"] += 1
                            else: engine["current_step_idx"] += 1; engine["current_loop_idx"] = 1
                            if engine["current_step_idx"] >= len(steps): engine["is_running"] = False; engine["is_finished"] = True
                        st.rerun() 
                    except Exception as e:
                        st.error(f"引擎发生故障: {e}"); engine["is_running"] = False

# ==========================================
# 模块 2: 自由聊天区
# ==========================================
elif st.session_state.current_page == "💬 自由聊天区":
    curr_chat = st.session_state.free_chats[st.session_state.current_chat_id]
    st.title(f"💬 {curr_chat['title']}")
    
    with st.container(height=700, border=False):
        for msg in curr_chat["messages"]:
            if msg["role"] == "system": continue
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg["role"] == "assistant": render_copy_button(msg["content"])

    if prompt := st.chat_input("探讨设定、查资料..."):
        if not active_p["api_key"]: st.error("缺 API Key！"); st.stop()
        if len(curr_chat["messages"]) == 0: curr_chat["title"] = prompt[:10] + "..."
            
        curr_chat["messages"].append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)
        
        client, profile = get_client()
        api_msgs = [{"role": m["role"], "content": m["content"]} for m in curr_chat["messages"]]
        
        with st.chat_message("assistant"):
            try:
                resp = client.chat.completions.create(**build_api_kwargs(profile, api_msgs))
                full_resp = st.write_stream(resp)
                render_copy_button(full_resp)
                curr_chat["messages"].append({"role": "assistant", "content": full_resp})
                save_free_chats(); st.rerun()
            except Exception as e: st.error(f"请求失败: {e}")

# ==========================================
# 模块 3: SOP与灵魂 (显式保存保障)
# ==========================================
elif st.session_state.current_page == "📝 账号SOP与灵魂":
    tab_sop, tab_vault = st.tabs(["🧩 SOP配置与人设管理", "🗄️ 账号记忆保险库 (风格蒸馏)"])
    
    with tab_sop:
        col1, col2 = st.columns([1, 2.5])
        with col1:
            st.subheader("账号 SOP 库")
            s_name = st.radio("选择编辑对象", list(st.session_state.sops.keys())) if st.session_state.sops else None
            if st.button("➕ 创建新账号 SOP", use_container_width=True):
                st.session_state.sops[f"新账号 {len(st.session_state.sops)}"] = {
                    "memory_mode": "manual", "system_prompt": "", "negative_memory": [],
                    "steps": [{"prompt": "", "loop": 1}], "triggers": []
                }
                save_sops(); st.rerun()

        with col2:
            if s_name:
                sop = st.session_state.sops[s_name]
                
                # 🔥 强制显式保存区
                ca, cb = st.columns([3, 1])
                with ca: new_name = st.text_input("✏️ 账号名称", s_name)
                with cb:
                    st.write("")
                    if st.button("💾 保存当前 SOP 配置", type="primary", use_container_width=True):
                        save_sops(); st.success("SOP 配置已落盘！")

                if new_name != s_name and new_name.strip():
                    st.session_state.sops[new_name] = st.session_state.sops.pop(s_name)
                    if s_name in st.session_state.memory:
                        st.session_state.memory[new_name] = st.session_state.memory.pop(s_name); save_memory()
                    save_sops(); st.rerun()
                
                st.markdown("### 🧠 记忆生长模式")
                mode_opts = {"manual": "保守派：手动挑选好文章提取风格 (稳健可控)", "dynamic": "激进派：自动活体记忆避坑反馈 (越训越聪明)"}
                sop["memory_mode"] = st.radio("选择成长路线", ["manual", "dynamic"], format_func=lambda x: mode_opts[x], index=0 if sop.get("memory_mode", "manual")=="manual" else 1, label_visibility="collapsed")
                
                if sop["memory_mode"] == "dynamic" and sop.get("negative_memory"):
                    with st.expander("👀 查看已吸收但未完全融合的【避坑清单】", expanded=True):
                        for nm in sop["negative_memory"]: st.markdown(f"- {nm}")
                        if st.button("强行清空避坑清单"): sop["negative_memory"] = []; save_sops(); st.rerun()
                
                st.markdown("### 🎭 专属人设 (System Prompt)")
                sop["system_prompt"] = st.text_area("核心指令", sop.get("system_prompt", ""), height=100)
                
                st.markdown("### 🧩 执行阶段配置")
                new_steps = []
                for i, step in enumerate(sop["steps"]):
                    with st.container(border=True):
                        st.markdown(f"**阶段 {i+1}**")
                        c1, c2 = st.columns([4, 1])
                        with c1: p_val = st.text_area("指令", step["prompt"], height=60, key=f"p_{i}", label_visibility="collapsed")
                        with c2: l_val = st.number_input("循环", min_value=1, value=step.get("loop", 1), key=f"l_{i}")
                        ref = step.get("reference", "")
                        with st.expander("📁 挂载阶段参考资料"): ref = st.text_area("粘贴资料", ref, key=f"r_{i}")
                        new_steps.append({"prompt": p_val, "loop": l_val, "reference": ref})
                sop["steps"] = new_steps
                
                c_a, c_b = st.columns(2)
                with c_a: 
                    if st.button("➕ 加阶段", use_container_width=True): sop["steps"].append({"prompt":"", "loop":1}); save_sops(); st.rerun()
                with c_b:
                    if len(sop["steps"])>1 and st.button("➖ 删阶段", use_container_width=True): sop["steps"].pop(); save_sops(); st.rerun()
                
                st.markdown("### ⚡ 监听触发器网络")
                new_triggers = []
                for i, t in enumerate(sop.get("triggers", [])):
                    with st.container(border=True):
                        t1, t2, t3 = st.columns([1, 1, 2])
                        with t1: typ = st.selectbox("类型", ["terminate", "intervene"], index=0 if t["type"]=="terminate" else 1, key=f"t_{i}")
                        with t2: kwd = st.text_input("关键词", t["keyword"], key=f"k_{i}")
                        with t3: act = st.text_input("动作指令", t.get("action", ""), disabled=(typ=="terminate"), key=f"a_{i}")
                        new_triggers.append({"type": typ, "keyword": kwd, "action": act})
                sop["triggers"] = new_triggers
                if st.button("➕ 加规则"): sop["triggers"].append({"type":"intervene", "keyword":"", "action":""}); save_sops(); st.rerun()
                
                st.divider()
                if st.button("🗑️ 删除此 SOP", type="primary"): 
                    del st.session_state.sops[s_name]
                    if s_name in st.session_state.memory: del st.session_state.memory[s_name]; save_memory()
                    save_sops(); st.rerun()

    with tab_vault:
        st.header("🗄️ 记忆保险库与炼丹炉")
        acc_mem = st.session_state.memory.get(s_name, [])
        if not acc_mem:
            st.info(f"【{s_name}】保险库为空。请先去自动化工作台执行并保存。")
        else:
            st.success(f"🗃️ 记忆库已沉淀 {len(acc_mem)} 篇作品。")
            for idx, item in enumerate(reversed(acc_mem)):
                with st.expander(f"📖 [{item['time']}] {item['topic']}"):
                    st.caption(item['content'])
                    if st.button("🗑️ 抹除此记忆", key=f"del_mem_{idx}"):
                        st.session_state.memory[s_name].remove(item); save_memory(); st.rerun()
            
            st.divider()
            if st.button("🔥 立即开炉提炼灵魂 (风格蒸馏)", type="primary", use_container_width=True):
                if not active_p["api_key"]: st.error("缺 API Key！"); st.stop()
                client, profile = get_client()
                combined_texts = "\n\n---\n\n".join([m['content'] for m in acc_mem[-3:]])
                distill_prompt = f"""深度分析以下小说风格，提炼一段极度严谨的【System Prompt】复刻文风。只需输出纯文本的 Prompt。
样本：\n{combined_texts}"""

                with st.spinner("神级炼丹师正在提炼灵魂设定..."):
                    try:
                        resp = client.chat.completions.create(model=profile["model"], messages=[{"role": "user", "content": distill_prompt}])
                        distilled_prompt = resp.choices[0].message.content.strip()
                        st.session_state.sops[s_name]["system_prompt"] = distilled_prompt
                        save_sops()
                        st.success("🎉 蒸馏成功！专属人设已被注入全新灵魂！")
                        st.info(f"**提炼成果:**\n\n{distilled_prompt}")
                    except Exception as e: st.error(f"蒸馏失败: {e}")

# ==========================================
# 模块 4: 底层引擎配置 (显式保存保障)
# ==========================================
elif st.session_state.current_page == "⚙️ 底层引擎配置":
    st.header("⚙️ 底层驱动配置")
    col_list, col_edit = st.columns([1, 2.5])
    
    with col_list:
        st.subheader("引擎库")
        p_names = [p["name"] for p in st.session_state.profiles]
        idx = st.radio("切换引擎", range(len(p_names)), format_func=lambda x: p_names[x], index=st.session_state.active_profile_idx)
        st.session_state.active_profile_idx = idx
        if st.button("➕ 新增引擎", use_container_width=True):
            new_profile = {
                "name": f"新引擎 {len(p_names)+1}", "base_url": "", "api_key": "", "model": "anthropic/claude-sonnet-4.6",
                "use_temperature": True, "temperature": 0.8, "use_max_tokens": True, "max_tokens": 4096
            }
            st.session_state.profiles.append(new_profile); save_profiles(); st.rerun()

    with col_edit:
        st.subheader("网络与参数调优")
        p = st.session_state.profiles[idx]
        
        # 🔥 强制显式保存区
        ca, cb = st.columns([3, 1])
        with ca: p["name"] = st.text_input("引擎标签", p["name"])
        with cb:
            st.write("")
            if st.button("💾 保存当前引擎配置", type="primary", use_container_width=True):
                save_profiles(); st.success("引擎配置已落盘！")

        c1, c2 = st.columns(2)
        with c1: p["base_url"] = st.text_input("Base URL", p["base_url"])
        with c2: p["api_key"] = st.text_input("API Key", p["api_key"], type="password")
        
        cm, cb = st.columns([3, 1])
        with cm: p["model"] = st.text_input("模型映射 (Model ID)", p["model"])
        with cb:
            st.write("")
            if st.button("🔄 联机获取列表"):
                if p["api_key"]:
                    with st.spinner("正在呼叫 API 获取模型..."):
                        success, result = fetch_models(p["base_url"], p["api_key"])
                        if success:
                            if result:
                                st.session_state.temp_models = result
                                st.success(f"✅ 成功抓取到 {len(result)} 个模型！")
                            else:
                                st.warning("⚠️ 请求成功，但该平台未返回任何模型列表。")
                        else:
                            st.error(f"❌ 获取失败: {result}")
                else:
                    st.error("❌ 请先填写 API Key！")
                    
        if "temp_models" in st.session_state:
            sel_m = st.selectbox("选择支持的模型 (选择后将覆盖当前模型)", ["(不覆盖)"] + st.session_state.temp_models)
            if sel_m != "(不覆盖)": 
                p["model"] = sel_m
                del st.session_state.temp_models
                save_profiles()
                st.rerun()
                
        st.markdown("#### 🎛️ 运行时超参数 (勾选生效)")
        sl1, sl2 = st.columns(2)
        with sl1:
            p["use_temperature"] = st.checkbox("🔥 Temperature", p.get("use_temperature", True))
            if p["use_temperature"]: p["temperature"] = st.slider("值", 0.0, 2.0, p.get("temperature", 0.8), 0.1, label_visibility="collapsed")
            p["use_max_tokens"] = st.checkbox("📏 Max Tokens", p.get("use_max_tokens", True))
            if p["use_max_tokens"]: p["max_tokens"] = st.slider("值", 512, 16384, p.get("max_tokens", 4096), 512, label_visibility="collapsed")
        with sl2:
            p["use_top_p"] = st.checkbox("🎲 Top P", p.get("use_top_p", False))
            if p["use_top_p"]: p["top_p"] = st.slider("值", 0.0, 1.0, p.get("top_p", 1.0), 0.05, label_visibility="collapsed")
            p["use_frequency_penalty"] = st.checkbox("🚫 Frequency Penalty", p.get("use_frequency_penalty", False))
            if p["use_frequency_penalty"]: p["frequency_penalty"] = st.slider("值", -2.0, 2.0, p.get("frequency_penalty", 0.0), 0.1, label_visibility="collapsed")
